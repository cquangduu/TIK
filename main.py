"""
================================================================================
DAILY KOREAN — main.py (CONTENT AUTOMATION SYSTEM)
================================================================================
Architecture:
    Phase 1  → Crawl News + Ra đề thi TOPIK 54
    Phase 2  → Viết văn mẫu + Phân tích từ vựng & ngữ pháp
    Phase 3  → Multi-channel editor → JSON cấu trúc cho 4 video TikTok + Word doc
    Phase 4  → Deep Dive Episode → JSON cấu trúc cho Video 5 (YouTube dài)
    Assets   → generate_tiktok_assets() → 5 video với segment-based MP3
    Render   → 5× remotion render (5 CompositionID khác nhau)
    Upload   → Google Drive (Word + 5 Video + YouTube metadata)
================================================================================
CHANGES (v3.0 — DAILY KOREAN Edition):
    - Brand name: DAILY KOREAN (데일리 코리안)
    - Azure TTS thay thế edge-tts (Chất lượng cao hơn)
    - Phase 4: Kịch bản Deep Dive Episode (~5-10 phút)
    - Video 5: YouTube Deep Dive với segment-based audio
    - YouTube Metadata: Auto-generate timestamps, title, hashtags
================================================================================
"""

import os
import json
import requests
import logging
import asyncio
import shutil
import random
import re
from bs4 import BeautifulSoup
from pydub import AudioSegment
import traceback
from datetime import datetime
import subprocess
import time
import platform

# ==================== AZURE TTS ====================
try:
    import azure.cognitiveservices.speech as speechsdk
    AZURE_TTS_AVAILABLE = True
except ImportError:
    AZURE_TTS_AVAILABLE = False
    logging.warning("⚠️ azure-cognitiveservices-speech not installed. Install with: pip install azure-cognitiveservices-speech")

# ==================== EDGE TTS FALLBACK ====================
try:
    import edge_tts
    EDGE_TTS_AVAILABLE = True
except ImportError:
    EDGE_TTS_AVAILABLE = False
    logging.warning("⚠️ edge-tts not installed. Install with: pip install edge-tts")

# Audio duration detection
try:
    from mutagen.mp3 import MP3
    MUTAGEN_AVAILABLE = True
except ImportError:
    MUTAGEN_AVAILABLE = False
    logging.warning("⚠️ mutagen not installed. Install with: pip install mutagen")

try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("⚠️ Thiếu thư viện python-docx. Hãy chạy: pip install python-docx")

# ==================== CONFIGURATION ====================
try:
    from dotenv import load_dotenv
    load_dotenv()
except ImportError:
    pass

# ==================== GOOGLE DRIVE UPLOAD ====================
import os.path
from google.auth.transport.requests import Request
from google.oauth2.credentials import Credentials
from google_auth_oauthlib.flow import InstalledAppFlow
from googleapiclient.discovery import build
from googleapiclient.http import MediaFileUpload

# ==================== YOUTUBE UPLOAD ====================
try:
    from youtube_uploader import (
        YouTubeUploader, 
        upload_tiktok_to_youtube, 
        upload_deep_dive_to_youtube
    )
    YOUTUBE_UPLOAD_AVAILABLE = True
except ImportError:
    YOUTUBE_UPLOAD_AVAILABLE = False
    logging.warning("⚠️ youtube_uploader module not found. YouTube upload disabled.")

# ==================== BLOG GENERATOR ====================
try:
    from blog_generator import generate_blog_from_data, BlogGenerator
    BLOG_GENERATOR_AVAILABLE = True
except ImportError:
    BLOG_GENERATOR_AVAILABLE = False
    logging.warning("⚠️ blog_generator module not found. Blog generation disabled.")

# ==================== PODCAST GENERATOR ====================
try:
    from podcast_generator import generate_podcast_from_data, PodcastGenerator
    PODCAST_GENERATOR_AVAILABLE = True
except ImportError:
    PODCAST_GENERATOR_AVAILABLE = False
    logging.warning("⚠️ podcast_generator module not found. Podcast generation disabled.")

# ==================== SOCIAL MEDIA PUBLISHER ====================
try:
    from social_publisher import publish_to_social_media, SocialMediaPublisher
    SOCIAL_PUBLISHER_AVAILABLE = True
except ImportError:
    SOCIAL_PUBLISHER_AVAILABLE = False
    logging.warning("⚠️ social_publisher module not found. Social media publishing disabled.")

# ==================== GITHUB DEPLOYER ====================
try:
    from github_deployer import deploy_blog_to_github, GitHubDeployer
    GITHUB_DEPLOYER_AVAILABLE = True
except ImportError:
    GITHUB_DEPLOYER_AVAILABLE = False
    logging.warning("⚠️ github_deployer module not found. GitHub deployment disabled.")

# ==================== MONETIZATION ====================
try:
    from monetization import MonetizationManager
    MONETIZATION_AVAILABLE = True
except ImportError:
    MONETIZATION_AVAILABLE = False
    logging.warning("⚠️ monetization module not found. Monetization features disabled.")

# ==================== TELEGRAM BOT ====================
try:
    from telegram_bot import send_daily_push
    TELEGRAM_BOT_AVAILABLE = True
except ImportError:
    TELEGRAM_BOT_AVAILABLE = False
    logging.warning("⚠️ telegram_bot module not found. Telegram push disabled.")


# ==================== GOOGLE DRIVE SCOPE ====================
GDRIVE_SCOPES = ['https://www.googleapis.com/auth/drive.file']
GDRIVE_TOKEN_FILE = 'drive_token.json'  # Separate token file for Drive

def upload_to_drive(file_path, folder_id):
    """Upload file lên Drive dùng drive_token.json hoặc tạo token mới"""
    if not os.path.exists(file_path):
        logging.error(f"❌ File không tồn tại để upload: {file_path}")
        return None

    logging.info(f"☁️  Đang upload lên Drive: {os.path.basename(file_path)}...")

    creds = None
    
    # Try to load existing Drive token
    if os.path.exists(GDRIVE_TOKEN_FILE):
        try:
            creds = Credentials.from_authorized_user_file(GDRIVE_TOKEN_FILE, GDRIVE_SCOPES)
        except Exception as e:
            logging.warning(f"⚠️ Token file invalid: {e}")
            creds = None
    
    # Refresh or get new credentials
    if not creds or not creds.valid:
        if creds and creds.expired and creds.refresh_token:
            try:
                from google.auth.transport.requests import Request
                creds.refresh(Request())
                logging.info("🔄 Refreshed Drive credentials")
            except Exception as e:
                logging.warning(f"⚠️ Could not refresh Drive token: {e}")
                creds = None
        
        if not creds:
            # Need to create new token - check for client_secrets.json
            if os.path.exists('client_secrets.json'):
                try:
                    from google_auth_oauthlib.flow import InstalledAppFlow
                    flow = InstalledAppFlow.from_client_secrets_file('client_secrets.json', GDRIVE_SCOPES)
                    creds = flow.run_local_server(port=8080)
                    logging.info("✅ New Drive credentials obtained")
                except Exception as e:
                    logging.error(f"❌ Could not create Drive credentials: {e}")
                    return None
            else:
                logging.error("❌ Không tìm thấy client_secrets.json! Không thể upload.")
                return None
        
        # Save credentials for next time
        if creds:
            with open(GDRIVE_TOKEN_FILE, 'w') as token:
                token.write(creds.to_json())
            logging.info(f"💾 Saved Drive credentials to {GDRIVE_TOKEN_FILE}")

    try:
        service = build('drive', 'v3', credentials=creds)
        file_metadata = {
            'name': os.path.basename(file_path),
            'parents': [folder_id]
        }
        media = MediaFileUpload(file_path, resumable=True)
        file = service.files().create(
            body=file_metadata,
            media_body=media,
            fields='id'
        ).execute()

        logging.info(f"✅ Upload thành công! File ID: {file.get('id')}")
        return file.get('id')

    except Exception as e:
        logging.error(f"❌ Lỗi khi upload lên Drive: {e}")
        return None


# ==================== DIRECTORY & ENV ====================
OUTPUT_DIR = "topik-video/public"
ASSETS_DIR = "topik-video/public/assets"
TEMP_DIR   = "temp_processing"

for _d in [OUTPUT_DIR, ASSETS_DIR, TEMP_DIR]:
    os.makedirs(_d, exist_ok=True)

GEMINI_API_KEY = os.getenv("GEMINI_API_KEY", "")
PEXELS_API_KEY = os.getenv("PEXELS_API_KEY", "")

# ==================== AZURE TTS CONFIGURATION ====================
AZURE_SPEECH_KEY = os.getenv("AZURE_SPEECH_KEY", "")
AZURE_SPEECH_REGION = os.getenv("AZURE_SPEECH_REGION", "eastasia")

# Voice assignment for different roles (Korean only)
AZURE_VOICE_CONFIG = {
    "host": "ko-KR-SunHiNeural",       # Dẫn chương trình & News (nữ, thân thiện)
    "news": "ko-KR-SunHiNeural",       # News Reader (nữ, healing vibes)
    "exam": "ko-KR-InJoonNeural",      # Đề thi (nam, nghiêm túc)
    "analysis": "ko-KR-JiMinNeural",   # Giải thích/Phân tích (nữ, chuyên nghiệp)
    "teaching": "ko-KR-InJoonNeural",  # Giảng dạy (nam, dứt khoát)
}

RSS_SOURCES = [
    {"name": "Donga Editorial",    "url": "https://rss.donga.com/editorial.xml"},
    {"name": "Hankyoreh Opinion",  "url": "http://www.hani.co.kr/rss/opinion/"},
    {"name": "MK News Editorial",  "url": "https://www.mk.co.kr/rss/30000001/"}
]

logging.basicConfig(level=logging.INFO, format='%(asctime)s [%(levelname)s] %(message)s')

# ==================== 5-VIDEO RENDER MANIFEST ====================
# Mapping: video_key → (CompositionID, output_filename_prefix)
VIDEO_MANIFEST = [
    {"key": "video_1_news",         "composition": "TikTok-NewsHealing",    "prefix": "V1_News",     "audio": "v1_news.mp3"},
    {"key": "video_2_outline",      "composition": "TikTok-WritingCoach",   "prefix": "V2_Writing",  "audio": "v2_outline.mp3"},
    {"key": "video_3_vocab_quiz",   "composition": "TikTok-VocabQuiz",      "prefix": "V3_Vocab",    "audio": "v3_vocab_quiz.mp3"},
    {"key": "video_4_grammar_quiz", "composition": "TikTok-GrammarQuiz",    "prefix": "V4_Grammar",  "audio": "v4_grammar_quiz.mp3"},
    {"key": "video_5_deep_dive",    "composition": "YouTube-DeepDive",      "prefix": "V5_DeepDive", "audio": "v5_deep_dive.mp3"},
]


# ==============================================================================
# 1. HELPER / UTILITY FUNCTIONS  (Giữ nguyên logic gốc)
# ==============================================================================

def call_ai_api(prompt, temperature=0.7):
    """Gọi Gemini API → trả về dict (JSON đã parse sạch)."""
    if not GEMINI_API_KEY:
        logging.error("❌ Chưa có GEMINI_API_KEY!")
        return {}

    url = (
        "https://generativelanguage.googleapis.com/v1beta/"
        f"models/gemini-2.5-flash:generateContent?key={GEMINI_API_KEY}"
    )
    headers = {'Content-Type': 'application/json'}
    data = {
        "contents": [{"parts": [{"text": prompt}]}],
        "generationConfig": {
            "temperature": temperature,
            "responseMimeType": "application/json"
        }
    }

    try:
        response = requests.post(url, headers=headers, json=data, timeout=200)
        if response.status_code != 200:
            logging.error(f"API Error: {response.text}")
            return {}

        result   = response.json()
        raw_text = result['candidates'][0]['content']['parts'][0]['text']

        # --- Làm sạch Markdown wrapper ---
        if "```json" in raw_text:
            raw_text = raw_text.split("```json")[1].split("```")[0]
        elif "```" in raw_text:
            raw_text = raw_text.split("```")[1].split("```")[0]

        clean_json = raw_text.strip()
        return json.loads(clean_json)

    except json.JSONDecodeError as e:
        logging.warning(f"⚠️  JSON lỗi nhẹ, đang thử sửa... ({e})")
        try:
            match = re.search(r"\{.*\}", raw_text, re.DOTALL)
            if match:
                return json.loads(match.group(0))
        except Exception:
            logging.error(f"❌ KHÔNG THỂ SỬA JSON. Raw: {raw_text[:200]}...")
            return {}
    except Exception as e:
        logging.error(f"❌ Lỗi hệ thống AI: {e}")
        return {}


def get_latest_editorial_rss():
    """Tìm bài editorial mới nhất từ RSS → trả về (url, source_name)."""
    logging.info("🔍 Đang tìm bài báo xã luận từ RSS feeds...")
    headers = {'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'}

    for source in RSS_SOURCES:
        try:
            logging.info(f"   Đang thử: {source['name']}...")
            response = requests.get(source['url'], headers=headers, timeout=10)

            try:
                soup = BeautifulSoup(response.content, 'xml')
            except Exception:
                soup = BeautifulSoup(response.content, 'html.parser')

            item = soup.find('item')
            if item:
                link = item.find('link').text.strip()
                title = item.find('title').text.strip()
                logging.info(f"✅ Tìm thấy: {title[:50]}... từ {source['name']}")
                return link, source['name']
        except Exception as e:
            logging.warning(f"   Lỗi với {source['name']}: {e}")
            continue

    return None, None


def extract_content(url_input):
    """Tải nội dung bài báo từ URL → trả về dict {title, text, url}."""
    if isinstance(url_input, (tuple, list)):
        url = url_input[0]
    else:
        url = url_input

    url = str(url).strip()
    logging.info(f"📥 Đang tải: {url}")

    headers = {
        'User-Agent': (
            'Mozilla/5.0 (Windows NT 10.0; Win64; x64) '
            'AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36'
        ),
    }

    try:
        response = requests.get(url, headers=headers, timeout=15)
        if response.encoding == 'ISO-8859-1':
            response.encoding = response.apparent_encoding

        soup = BeautifulSoup(response.text, 'html.parser')

        selectors = ['div.text', '#article_view', '.article_txt', '.art_txt', 'article', '#news_body_id']
        main_div = None
        for sel in selectors:
            main_div = soup.select_one(sel)
            if main_div:
                break

        if main_div:
            paragraphs = main_div.find_all('p')
            text_list = [p.get_text().strip() for p in paragraphs if len(p.get_text()) > 20]
        else:
            paragraphs = soup.find_all('p')
            text_list = [p.get_text().strip() for p in paragraphs if len(p.get_text()) > 50]

        article_text = "\n".join(text_list)

        if len(article_text) < 200:
            logging.warning("⚠️  Nội dung quá ngắn.")
            return None

        title = soup.title.string if soup.title else "News"
        return {'title': title, 'text': article_text, 'url': url}

    except Exception as e:
        logging.error(f"❌ Lỗi tải bài: {e}")
        return None


def get_video_duration(file_path: str) -> float:
    """
    Get video duration in seconds using ffprobe (if available) or Pexels API data.
    
    Args:
        file_path: Path to the video file
        
    Returns:
        Duration in seconds (float), or 0.0 if unable to detect
    """
    if not os.path.exists(file_path):
        logging.warning(f"⚠️ Video file not found: {file_path}")
        return 0.0
    
    # Try using ffprobe (comes with ffmpeg)
    try:
        result = subprocess.run(
            [
                "ffprobe", "-v", "error", 
                "-show_entries", "format=duration",
                "-of", "default=noprint_wrappers=1:nokey=1",
                file_path
            ],
            capture_output=True, text=True, timeout=30
        )
        if result.returncode == 0 and result.stdout.strip():
            duration = float(result.stdout.strip())
            logging.info(f"📹 Video duration: {duration:.2f}s (from ffprobe)")
            return duration
    except (subprocess.TimeoutExpired, FileNotFoundError, ValueError) as e:
        logging.debug(f"ffprobe not available or failed: {e}")
    
    # Fallback: estimate from file size (rough approximation)
    # Average TikTok video: ~2MB per 10 seconds at 1080p
    try:
        file_size_mb = os.path.getsize(file_path) / (1024 * 1024)
        estimated_duration = file_size_mb / 0.2  # ~0.2 MB per second rough estimate
        logging.info(f"📹 Video duration: ~{estimated_duration:.1f}s (estimated from file size)")
        return estimated_duration
    except Exception:
        pass
    
    # If all else fails, return 0
    logging.warning(f"⚠️ Could not determine video duration for: {file_path}")
    return 0.0


# Global variable to store video background duration
VIDEO_BG_DURATION_CACHE = 0.0


def download_background_video(query, output_path):
    """
    Tải video nền từ Pexels API và lưu duration.
    
    Returns:
        dict: {"success": bool, "duration": float} or False for backward compatibility
    """
    global VIDEO_BG_DURATION_CACHE
    
    if not PEXELS_API_KEY:
        logging.warning("⚠️  Thiếu PEXELS_API_KEY.")
        return False

    clean_query = "".join(e for e in query if e.isalnum() or e.isspace())
    logging.info(f"🎬 Đang tìm video nền: '{clean_query}'...")

    headers = {"Authorization": PEXELS_API_KEY}
    api_url = (
        f"https://api.pexels.com/videos/search?"
        f"query={clean_query}&per_page=1&orientation=portrait"
    )

    try:
        response = requests.get(api_url, headers=headers, timeout=15)
        data = response.json()
        if not data.get('videos'):
            return False

        video_data  = random.choice(data['videos'])
        video_files = video_data['video_files']
        valid_files = [v for v in video_files if v['width'] and 600 < v['width'] < 1200]
        best_file   = valid_files[0] if valid_files else video_files[0]
        
        # Get duration from Pexels API (available in video_data)
        pexels_duration = video_data.get('duration', 0)

        with requests.get(best_file['link'], stream=True) as r:
            r.raise_for_status()
            with open(output_path, 'wb') as f:
                for chunk in r.iter_content(chunk_size=8192):
                    f.write(chunk)

        # Get actual duration from file (more accurate than Pexels API)
        actual_duration = get_video_duration(output_path)
        if actual_duration <= 0:
            actual_duration = pexels_duration  # Fallback to Pexels API duration
        
        VIDEO_BG_DURATION_CACHE = actual_duration
        logging.info(f"✅ Đã tải video nền! Duration: {actual_duration:.2f}s")
        return {"success": True, "duration": actual_duration}

    except Exception as e:
        logging.error(f"❌ Lỗi tải video: {e}")
        return False


def sanitize_text(text):
    """Lọc ký tự lỗi XML để tránh crash file Word."""
    if not text:
        return ""
    return re.sub(r'[^\x09\x0A\x0D\x20-\uD7FF\uE000-\uFFFD]', '', str(text))


# ==============================================================================
# 2. AZURE TTS FUNCTIONS  —  Thay thế edge-tts (Chất lượng cao hơn)
# ==============================================================================

# ═══════════════════════════════════════════════════════════════════════════════
# SSML Dynamic Rate Configuration
# ═══════════════════════════════════════════════════════════════════════════════
# Ngưỡng ký tự để điều chỉnh tốc độ đọc tự động
# Văn bản dài hơn sẽ được đọc nhanh hơn để khớp với khung hình video
SSML_RATE_THRESHOLDS = {
    "short": {"max_chars": 50, "rate": "+0%"},       # Ngắn: tốc độ bình thường
    "medium": {"max_chars": 150, "rate": "+10%"},   # Trung bình: tăng 10%
    "long": {"max_chars": 300, "rate": "+15%"},     # Dài: tăng 15%
    "very_long": {"max_chars": 500, "rate": "+20%"},# Rất dài: tăng 20%
    "extra_long": {"max_chars": float('inf'), "rate": "+25%"}  # Siêu dài: tăng 25%
}

# Target max duration for TikTok videos (seconds)
TIKTOK_MAX_DURATION = 55  # Target < 60s, aim for 55s
TIKTOK_COMPRESS_RATE = "+15%"  # Rate to apply when total > 55s


def estimate_reading_time(text: str, chars_per_second: float = 5.0) -> float:
    """
    Estimate reading time in seconds based on text length.
    Korean typically reads at ~4-6 characters per second.
    
    Args:
        text: Korean text
        chars_per_second: Reading speed (default 5.0 for Korean)
        
    Returns:
        Estimated duration in seconds
    """
    if not text:
        return 0.0
    return len(text.strip()) / chars_per_second


def should_compress_audio(total_text: str, target_max: float = TIKTOK_MAX_DURATION) -> tuple[bool, str]:
    """
    Check if audio should be compressed based on estimated duration.
    
    If estimated reading time > target_max seconds, return True with +15% rate.
    
    Args:
        total_text: Combined Korean text for the video
        target_max: Maximum allowed duration in seconds
        
    Returns:
        Tuple of (should_compress: bool, rate: str)
    """
    estimated = estimate_reading_time(total_text)
    
    if estimated > target_max:
        logging.info(f"⚡ Auto-compress: Estimated {estimated:.1f}s > {target_max}s → applying {TIKTOK_COMPRESS_RATE}")
        return True, TIKTOK_COMPRESS_RATE
    else:
        return False, "+0%"


def _calculate_dynamic_rate(text: str, base_rate: str = "+0%") -> str:
    """
    Calculate dynamic speech rate based on text length.
    
    Văn bản dài hơn sẽ được đọc nhanh hơn để nén thời gian đọc,
    giúp khớp với khung hình video mà vẫn giữ chất lượng tự nhiên.
    
    Args:
        text: Korean text to synthesize
        base_rate: Base rate specified by caller (e.g., "-10%", "+5%")
        
    Returns:
        Final rate string (e.g., "+15%")
    """
    text_length = len(text.strip())
    
    # Determine rate based on text length
    dynamic_rate_value = 0
    for tier_name, config in SSML_RATE_THRESHOLDS.items():
        if text_length <= config["max_chars"]:
            rate_str = config["rate"]
            dynamic_rate_value = int(rate_str.replace("%", "").replace("+", ""))
            break
    
    # Parse base rate
    base_rate_clean = base_rate.replace("%", "")
    if base_rate_clean.startswith("+"):
        base_rate_value = int(base_rate_clean[1:])
    elif base_rate_clean.startswith("-"):
        base_rate_value = int(base_rate_clean)
    else:
        base_rate_value = int(base_rate_clean) if base_rate_clean else 0
    
    # Combine rates (additive)
    final_rate_value = base_rate_value + dynamic_rate_value
    
    # Clamp to reasonable range (-50% to +50%)
    final_rate_value = max(-50, min(50, final_rate_value))
    
    # Format as string
    if final_rate_value >= 0:
        return f"+{final_rate_value}%"
    else:
        return f"{final_rate_value}%"


def _build_ssml(text: str, voice_name: str, rate: str, use_dynamic_rate: bool = True) -> str:
    """
    Build SSML markup with prosody rate adjustment.
    
    Args:
        text: Korean text to synthesize
        voice_name: Azure voice name
        rate: Base rate (will be adjusted dynamically if use_dynamic_rate=True)
        use_dynamic_rate: Whether to apply dynamic rate based on text length
        
    Returns:
        Complete SSML string
    """
    # Calculate final rate
    if use_dynamic_rate:
        final_rate = _calculate_dynamic_rate(text, rate)
    else:
        final_rate = rate
    
    # Clean up rate format for SSML
    rate_value = final_rate.replace("%", "")
    if not rate_value.startswith("+") and not rate_value.startswith("-"):
        rate_value = f"+{rate_value}"
    
    # Escape special XML characters in text
    escaped_text = (text
        .replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
        .replace('"', "&quot;")
        .replace("'", "&apos;")
    )
    
    # Build SSML with breaks for natural pauses
    ssml = f"""<speak version="1.0" xmlns="http://www.w3.org/2001/10/synthesis" xml:lang="ko-KR">
    <voice name="{voice_name}">
        <prosody rate="{rate_value}%">
            {escaped_text}
        </prosody>
    </voice>
</speak>"""
    
    return ssml


def generate_azure_tts(text: str, voice_name: str, output_path: str, rate: str = "+0%", use_dynamic_rate: bool = True) -> float:
    """
    Generate TTS audio using Azure Cognitive Services Speech SDK.
    
    Uses SSML with dynamic prosody rate adjustment:
    - Văn bản ngắn (< 50 chars): tốc độ bình thường
    - Văn bản trung bình (50-150 chars): +10%
    - Văn bản dài (150-300 chars): +15%
    - Văn bản rất dài (300-500 chars): +20%
    - Văn bản siêu dài (> 500 chars): +25%
    
    Args:
        text: Korean text to synthesize (KHÔNG tạo audio cho tiếng Việt!)
        voice_name: Azure voice name (e.g., "ko-KR-SunHiNeural")
        output_path: Path to save the MP3 file
        rate: Base speed rate (e.g., "-10%", "+0%", "+10%")
        use_dynamic_rate: Whether to apply dynamic rate based on text length
        
    Returns:
        Duration in seconds (float), or 0.0 if failed
        
    RULE: CHỈ tạo audio cho tiếng Hàn. Tiếng Việt dùng làm phụ đề, không có audio.
    """
    if not text or not text.strip():
        return 0.0
    
    # Check for Vietnamese text and REMOVE Vietnamese portions instead of skipping entirely
    # This handles cases where explanation_ko contains mixed Korean/Vietnamese
    vietnamese_pattern = r'[àáảãạăằắẳẵặâầấẩẫậèéẻẽẹêềếểễệìíỉĩịòóỏõọôồốổỗộơờớởỡợùúủũụưừứửữựỳýỷỹỵđ]'
    if re.search(vietnamese_pattern, text.lower()):
        # Remove Vietnamese portions: text in single quotes like 'nguyên nhân chính của việc...'
        # Also remove explanations in parentheses containing Vietnamese
        cleaned_text = re.sub(r"'[^']*[àáảãạăằắẳẵặâầấẩẫậèéẻẽẹêềếểễệìíỉĩịòóỏõọôồốổỗộơờớởỡợùúủũụưừứửữựỳýỷỹỵđ][^']*'", "", text)
        cleaned_text = re.sub(r"\([^)]*[àáảãạăằắẳẵặâầấẩẫậèéẻẽẹêềếểễệìíỉĩịòóỏõọôồốổỗộơờớởỡợùúủũụưừứửữựỳýỷỹỵđ][^)]*\)", "", cleaned_text)
        # Clean up multiple spaces and orphaned punctuation
        cleaned_text = re.sub(r'\s+', ' ', cleaned_text).strip()
        cleaned_text = re.sub(r'\s+([.,])', r'\1', cleaned_text)
        
        if cleaned_text and len(cleaned_text) > 10:
            logging.info(f"🔄 Removed Vietnamese from TTS text: '{text[:50]}...' → '{cleaned_text[:50]}...'")
            text = cleaned_text
        else:
            logging.warning(f"⚠️ Text mostly Vietnamese, skipping TTS: {text[:50]}...")
            return 0.0
    
    if not AZURE_TTS_AVAILABLE or not AZURE_SPEECH_KEY:
        logging.warning("⚠️ Azure TTS not available, falling back to edge-tts...")
        final_rate = _calculate_dynamic_rate(text, rate) if use_dynamic_rate else rate
        return _fallback_edge_tts_sync(text, voice_name, output_path, final_rate)
    
    try:
        # Configure Azure Speech SDK
        speech_config = speechsdk.SpeechConfig(
            subscription=AZURE_SPEECH_KEY,
            region=AZURE_SPEECH_REGION
        )
        speech_config.set_speech_synthesis_output_format(
            speechsdk.SpeechSynthesisOutputFormat.Audio16Khz32KBitRateMonoMp3
        )
        speech_config.speech_synthesis_voice_name = voice_name
        
        # Create audio config for file output
        audio_config = speechsdk.audio.AudioOutputConfig(filename=output_path)
        
        # Create synthesizer
        synthesizer = speechsdk.SpeechSynthesizer(
            speech_config=speech_config,
            audio_config=audio_config
        )
        
        # Build SSML with dynamic rate adjustment
        ssml = _build_ssml(text, voice_name, rate, use_dynamic_rate)
        
        # Log rate info for debugging
        if use_dynamic_rate:
            final_rate = _calculate_dynamic_rate(text, rate)
            logging.debug(f"📢 SSML rate: {final_rate} (text length: {len(text)} chars)")
        
        # Synthesize
        result = synthesizer.speak_ssml_async(ssml).get()
        
        if result.reason == speechsdk.ResultReason.SynthesizingAudioCompleted:
            duration = get_audio_duration(output_path)
            logging.debug(f"✅ Azure TTS OK: {os.path.basename(output_path)} ({duration:.2f}s)")
            return duration
        elif result.reason == speechsdk.ResultReason.Canceled:
            cancellation = result.cancellation_details
            logging.error(f"❌ Azure TTS canceled: {cancellation.reason}")
            if cancellation.reason == speechsdk.CancellationReason.Error:
                logging.error(f"   Error details: {cancellation.error_details}")
            # Fallback to edge-tts
            final_rate = _calculate_dynamic_rate(text, rate) if use_dynamic_rate else rate
            return _fallback_edge_tts_sync(text, voice_name, output_path, final_rate)
        else:
            logging.error(f"❌ Azure TTS failed with reason: {result.reason}")
            final_rate = _calculate_dynamic_rate(text, rate) if use_dynamic_rate else rate
            return _fallback_edge_tts_sync(text, voice_name, output_path, final_rate)
            
    except Exception as e:
        logging.error(f"❌ Azure TTS exception: {e}")
        final_rate = _calculate_dynamic_rate(text, rate) if use_dynamic_rate else rate
        return _fallback_edge_tts_sync(text, voice_name, output_path, final_rate)


def _fallback_edge_tts_sync(text: str, voice_name: str, output_path: str, rate: str) -> float:
    """Fallback to edge-tts when Azure TTS fails (synchronous wrapper)."""
    if not EDGE_TTS_AVAILABLE:
        logging.error("❌ Neither Azure TTS nor edge-tts available!")
        return 0.0
    
    try:
        asyncio.run(_tts_to_file(text, voice_name, rate, output_path))
        return get_audio_duration(output_path)
    except Exception as e:
        logging.error(f"❌ edge-tts fallback failed: {e}")
        return 0.0


async def generate_azure_tts_async(text: str, voice_name: str, output_path: str, rate: str = "+0%", use_dynamic_rate: bool = True) -> float:
    """
    Async version of generate_azure_tts for use in async contexts.
    Wraps the sync function in an executor.
    
    Args:
        text: Korean text to synthesize
        voice_name: Azure voice name
        output_path: Path to save the MP3 file
        rate: Base speed rate
        use_dynamic_rate: Whether to apply dynamic rate based on text length
    """
    import concurrent.futures
    from functools import partial
    
    loop = asyncio.get_event_loop()
    with concurrent.futures.ThreadPoolExecutor() as executor:
        # Use partial to pass all arguments including use_dynamic_rate
        func = partial(generate_azure_tts, text, voice_name, output_path, rate, use_dynamic_rate)
        duration = await loop.run_in_executor(executor, func)
    return duration


# ==============================================================================
# 2. AI PIPELINE  —  Phase 1 → Phase 2 → Phase 3
# ==============================================================================

def run_phase_1(article_text: str) -> dict:
    """
    Phase 1: Phân tích bài báo → ra đề thi TOPIK 54 + tóm tắt tin tức.
    Giữ nguyên prompt gốc.
    """
    logging.info("🧠 Phase 1: Phân tích & Ra đề...")

    prompt_p1 = f"""
    Bạn là chuyên gia ra đề thi TOPIK II với hơn 10 năm kinh nghiệm.

    Dựa trên thông tin sau (được trích từ một bản tin xã hội, KHÔNG phải đề thi):

    [NEWS_SUMMARY]
    {article_text[:3000]}

    Hãy thực hiện các nhiệm vụ sau:

    1. Phân tích và trích xuất MỘT "vấn đề xã hội hoặc xu hướng xã hội" mang tính tổng quát:
       - Không nhắc đến thời gian cụ thể
       - Không nhắc đến sự kiện hay tên tổ chức cụ thể
       - Không dùng văn phong báo chí
       - Phải là vấn đề có thể lặp lại nhiều năm

    2. Tạo MỘT đề thi viết TOPIK II – câu 54:
       - Đúng văn phong đề thi TOPIK
       - Dạng thảo luận: nguyên nhân → tác động → giải pháp
       - Không dùng số liệu chi tiết
       - Không quá thời sự

    3. Tóm tắt cơ bản nội dung bài báo bằng TIẾNG HÀN ĐƠN GIẢN (Level TOPIK 3) để làm bản tin dễ nghe.

    ⚠️ LƯU Ý QUAN TRỌNG:
       - Tuyệt đối KHÔNG nhắc đến nguồn tin, báo chí hay tên tổ chức
       - Phải khiến người học cảm giác đây là đề thi thật đã từng ra
       - Văn phong phải giống 90% đề TOPIK thật

    📋 CẤU TRÚC ĐỀ BÀI CHUẨN TOPIK 54:

    다음을 주제로 하여 자신의 생각을 600~700자로 글을 쓰십시오. (30점)

    [2-3 câu mở đầu giới thiệu vấn đề xã hội, không nhắc nguồn]

    <조건>
    • [Nguyên nhân của vấn đề]
    • [Tác động/ảnh hưởng của vấn đề]
    • [Giải pháp hoặc hướng đi tương lai]

    OUTPUT JSON (STRICT FORMAT):
    {{
        "topic_korean": "Chủ đề chính (tiếng Hàn)",
        "video_keyword": "từ khóa tìm video nền",
        "news_summary_easy_kr": "Tóm tắt tin tức thành chủ đề đơn giản (Tiếng Hàn)",
        "question_full_text": "Đề thi TOPIK 54 đầy đủ (Giữ nguyên độ khó cao cấp)"
    }}
    """

    data_p1 = call_ai_api(prompt_p1, temperature=0.5)
    if not data_p1:
        logging.error("❌ Phase 1 thất bại — không có dữ liệu.")
        return {}

    logging.info(f"🔹 Chủ đề: {data_p1.get('topic_korean', 'N/A')}")
    return data_p1


def run_phase_2(data_p1: dict) -> dict:
    """
    Phase 2: Viết văn mẫu + phân tích từ vựng & ngữ pháp.
    Giữ nguyên prompt gốc.
    """
    logging.info("🧠 Phase 2: Viết văn mẫu...")

    prompt_p2 = f"""
    Role: You are the Head Grader of the TOPIK II Writing Section (쓰기 채점 위원장).
    Input Question & Conditions: {data_p1['question_full_text']}

    OBJECTIVE: Write a model essay (모범 답안) that receives a PERFECT SCORE (50/50).

    --- 🛑 STRICT WRITING RULES (DO NOT IGNORE) ---
    1.  **FORMAT**:
        -   Total length: 600-700 characters (No more, no less).
        -   Structure: EXACTLY 4 Paragraphs.
            -   Para 1: Introduction (Generalize the issue).
            -   Para 2: Response to Bullet Point 1.
            -   Para 3: Response to Bullet Point 2.
            -   Para 4: Response to Bullet Point 3 + Conclusion.

    2.  **TONE & STYLE**:
        -   **Video Style**: Highly Academic & Formal (학술적 글쓰기).
        -   **Ending**: STRICTLY use '-다/는다' style. (NEVER use -습니다/해요).
        -   **Perspective**: Objective. Do NOT use "I" (저/나/제 생각에는). Use "It is considered that..." (여겨진다) or "We" (우리).
        -   **Banned Phrases**: '것 같다' (seems like), '알 수 있다' (can know). -> REPLACE WITH: '분명하다' (clear), '파악된다' (identified).

    3.  **ADVANCED VOCABULARY INJECTION**:
        -   Must use at least 2 **Four-character Idioms (사자성어)** (e.g., 설상가상, 역지사지, 유비무환...).
        -   Must use **Advanced Connectors**:
            -   Start Para 2 with: 우선 / 무엇보다도...
            -   Start Para 3 with: 하지만 / 반면에 / 이와 더불어...
            -   Start Para 4 with: 따라서 / 결론적으로...

    --- TASK LIST ---

    TASK 1: Write the Essay.

    TASK 2: Vocabulary Extraction (Korean -> Vietnamese).
    (Tuyệt đối không thêm Hán tự trong ngoặc đơn ở phần "word")
    -   Select all "Tier 1" academic words from the essay (e.g., phenomenon, implementation, countermeasure).

    TASK 3: Grammar Analysis.
    -   Explain all advanced grammar points used.

    OUTPUT JSON STRUCTURE:
    {{
        "essay": "Korean text...",
        "analysis_list": [
            {{ "item": "All vocabulary advanced from the essay", "professor_explanation": "Lời giảng của giáo sư về sắc thái/cách dùng..." }},
            {{ "item": "All grammar points from the essay", "professor_explanation": "..." }}
        ]
    }}
    """

    data_p2 = call_ai_api(prompt_p2, temperature=0.7)
    if not data_p2:
        logging.error("❌ Phase 2 thất bại — không có dữ liệu.")
        return {}

    logging.info(f"🔹 Essay length: {len(data_p2.get('essay', ''))} chars")
    return data_p2



def run_phase_3(data_p1: dict, data_p2: dict) -> dict:
    """
    Phase 3 — BIÊN TẬP VIÊN ĐA KÊNH (Multi-Channel Editor).

    Input:  data_p1 (News + Đề thi), data_p2 (Essay + Analysis)
    Output: JSON cấu trúc cho 4 video TikTok + dữ liệu Word doc.

    RULE: Korean Audio - Vietnamese Subtitles
    NEW: Mỗi video có opening_ment (lời chào tiếng Hàn)
    """
    logging.info("🧠 Phase 3: Biên tập viên đa kênh — cấu trúc 4 video TikTok...")

    analysis_str = json.dumps(data_p2.get('analysis_list', []), ensure_ascii=False)

    prompt_p3 = f"""
    Bạn là "Biên tập viên đa kênh" (Multi-Channel Content Editor) chuyên tạo nội dung học tiếng Hàn cho TikTok.

    ⛔ NGUYÊN TẮC CỐT LÕI:
    1. CHỈ sử dụng thông tin từ DỮ LIỆU ĐẦU VÀO dưới đây. Không tự ý sáng tạo nội dung mới.
    2. Không thay đổi nội dung bài văn mẫu hay đề thi gốc.
    3. Các câu hỏi trắc nghiệm (Quiz) phải dựa đúng vào từ vựng/ngữ pháp đã có trong 'Phân tích của Giáo sư'.
    4. Mỗi script phải ĐỦ NGẮN để đọc trong 30–45 giây (phù hợp TikTok).

    🎯 LUẬT BẮT BUỘC: KOREAN AUDIO - VIETNAMESE SUBTITLES
    - Tất cả audio sẽ được đọc bằng TIẾNG HÀN.
    - Tiếng Việt CHỈ dùng làm phụ đề hiển thị trên màn hình.
    - Mỗi video phải tách rõ: "audio_text" (Hàn) và "segments" (cặp ko/vi).

    🎤 KỊCH BẢN LỜI CHÀO (OPENING_MENT) + KẾT THÚC (CLOSING_MENT) - LIÊN KẾT 4 VIDEO:
    
    ⚡ QUAN TRỌNG: 4 video short phải LIÊN KẾT với nhau như một series!
    - Video 1 → mở đầu series hôm nay, kết thúc dẫn sang Video 2
    - Video 2 → nối tiếp từ Video 1, kết thúc dẫn sang Video 3  
    - Video 3 → nối tiếp từ Video 2, kết thúc dẫn sang Video 4
    - Video 4 → nối tiếp từ Video 3, kết thúc khép lại series hôm nay
    
    📌 OPENING_MENT (mở đầu ~3-5 giây):
    - Video 1: "안녕하세요! 데일리 코리안입니다. 오늘의 한국 사회 이슈, 함께 들어볼까요?"
    - Video 2: "자, 이어서! 방금 배운 주제로 토픽 쓰기 연습을 해볼까요?"
    - Video 3: "이제 어휘 퀴즈 시간입니다! 준비됐나요?"
    - Video 4: "마지막으로 문법 퀴즈입니다! 한번 도전해볼까요?"
    
    📌 CLOSING_MENT (kết thúc ~3-5 giây):
    - Video 1: "다음 영상에서 쓰기 연습 함께 해봐요!"
    - Video 2: "다음 영상에서 어휘 퀴즈 풀어봐요!"
    - Video 3: "다음 영상에서 문법 퀴즈도 도전해봐요!"
    - Video 4: "오늘도 수고했어요! 내일 또 만나요, 안녕!"

    --- DỮ LIỆU ĐẦU VÀO (SOURCE DATA) ---

    1. [TÓM TẮT TIN TỨC (Tiếng Hàn đơn giản)]:
    {data_p1.get('news_summary_easy_kr', '')}

    2. [ĐỀ THI GỐC TOPIK 54]:
    {data_p1.get('question_full_text', '')}

    3. [VĂN MẪU CHUẨN (4 đoạn)]:
    {data_p2.get('essay', '')}

    4. [PHÂN TÍCH CỦA GIÁO SƯ (Từ vựng + Ngữ pháp)]:
    {analysis_str}

    --- YÊU CẦU OUTPUT ---
    Trả về 1 JSON duy nhất với cấu trúc CHÍNH XÁC sau:

    {{
        "meta": {{
            "topic_title_vi": "Tiêu đề tiếng Việt hấp dẫn (dưới 10 từ)"
        }},

        "word_doc_data": {{
            "vocab_list": [
                {{"word": "từ vựng", "meaning_vi": "nghĩa tiếng Việt", "example": "1 câu ví dụ ngắn chứa từ này"}}
            ],
            "grammar_list": [
                {{"point": "điểm ngữ pháp", "meaning_vi": "nghĩa", "example": "1 câu ví dụ ngắn"}}
            ],
            "cloze_test": {{
                "question": "Trích 1 câu hay nhất từ [VĂN MẪU CHUẨN], thay từ khóa bằng [ ___ ]",
                "answer": "từ khóa bị che",
                "hint_vi": "gợi ý nghĩa tiếng Việt"
            }}
        }},

        "tiktok_script": {{
            "video_1_news": {{
                // Giọng nữ nhẹ nhàng — Healing vibes
                // Viết lại [TÓM TẮT TIN TỨC] theo phong cách thủ thỉ, tâm tình.
                // Kết bằng câu hỏi gợi mở: "Nếu thi vào chủ đề này thì sao?"
                "opening_ment": "안녕하세요! 데일리 코리안입니다. 오늘의 한국 사회 이슈, 함께 들어볼까요?",
                "audio_text": "Toàn bộ text tiếng Hàn để TTS đọc (20-30 giây, KHÔNG bao gồm opening_ment)",
                "closing_ment": "다음 영상에서 쓰기 연습 함께 해봐요!",
                "segments": [
                    {{"ko": "Câu tiếng Hàn 1", "vi": "Dịch tiếng Việt 1"}},
                    {{"ko": "Câu tiếng Hàn 2", "vi": "Dịch tiếng Việt 2"}},
                    {{"ko": "Câu tiếng Hàn 3", "vi": "Dịch tiếng Việt 3"}}
                ]
            }},
            "video_2_outline": {{
                // Giọng nam giáo sư — Teaching vibes
                // Tóm tắt cấu trúc bài văn mẫu thành 4 phần.
                "opening_ment": "자, 이어서! 방금 배운 주제로 토픽 쓰기 연습을 해볼까요?",
                "audio_text": "Toàn bộ text tiếng Hàn để TTS đọc (25-35 giây, KHÔNG bao gồm opening_ment)",
                "closing_ment": "다음 영상에서 어휘 퀴즈 풀어봐요!",
                "parts": [
                    {{
                        "role": "intro",
                        "label_vi": "Mở Bài",
                        "ko": "Luận điểm tiếng Hàn cho phần mở bài",
                        "vi": "Giải thích tiếng Việt ngắn gọn"
                    }},
                    {{
                        "role": "body_1",
                        "label_vi": "Thân Bài 1 - Nguyên nhân",
                        "ko": "Luận điểm tiếng Hàn",
                        "vi": "Giải thích tiếng Việt"
                    }},
                    {{
                        "role": "body_2",
                        "label_vi": "Thân Bài 2 - Tác động",
                        "ko": "Luận điểm tiếng Hàn",
                        "vi": "Giải thích tiếng Việt"
                    }},
                    {{
                        "role": "body_3",
                        "label_vi": "Kết Bài - Giải pháp",
                        "ko": "Luận điểm tiếng Hàn",
                        "vi": "Giải thích tiếng Việt"
                    }}
                ]
            }},
            "video_3_vocab_quiz": {{
                // Game show: Đọc câu hỏi → Im lặng 4s → Đáp án + Giải thích
                // Chọn 1 từ KHÓ NHẤT từ [PHÂN TÍCH CỦA GIÁO SƯ] để hỏi.
                "opening_ment": "이제 어휘 퀴즈 시간입니다! 준비됐나요?",
                "target_word": "từ vựng được chọn để hỏi",
                "question_ko": "Câu hỏi trắc nghiệm bằng tiếng Hàn (~5 giây đọc)",
                "question_vi": "Dịch câu hỏi sang tiếng Việt",
                "options_ko": ["A. Đáp án Hàn 1", "B. Đáp án Hàn 2", "C. Đáp án Hàn 3", "D. Đáp án Hàn 4"],
                "options_vi": ["A. Dịch Việt 1", "B. Dịch Việt 2", "C. Dịch Việt 3", "D. Dịch Việt 4"],
                "correct_answer": "C",
                "explanation_ko": "Giải thích tiếng Hàn ngắn gọn (~8 giây đọc)",
                "explanation_vi": "Giải thích tiếng Việt chi tiết hơn",
                "closing_ment": "다음 영상에서 문법 퀴즈도 도전해봐요!"
            }},
            "video_4_grammar_quiz": {{
                // Game show: Đọc câu hỏi → Im lặng 4s → Đáp án + Giải thích
                // Chọn 1 điểm ngữ pháp HAY NHẤT từ [PHÂN TÍCH CỦA GIÁO SƯ].
                "opening_ment": "마지막으로 문법 퀴즈입니다! 한번 도전해볼까요?",
                "target_grammar": "điểm ngữ pháp được chọn để hỏi",
                "question_ko": "Câu hỏi điền ngữ pháp vào chỗ trống bằng tiếng Hàn (~5 giây đọc)",
                "question_vi": "Dịch câu hỏi sang tiếng Việt",
                "options_ko": ["A. Đáp án Hàn 1", "B. Đáp án Hàn 2", "C. Đáp án Hàn 3", "D. Đáp án Hàn 4"],
                "options_vi": ["A. Dịch Việt 1", "B. Dịch Việt 2", "C. Dịch Việt 3", "D. Dịch Việt 4"],
                "correct_answer": "C",
                "explanation_ko": "Giải thích tiếng Hàn ngắn gọn (~8 giây đọc)",
                "explanation_vi": "Giải thích tiếng Việt chi tiết hơn",
                "closing_ment": "오늘도 수고했어요! 내일 또 만나요, 안녕!"
            }}
        }}
    }}
    """

    data_p3 = call_ai_api(prompt_p3, temperature=0.7)
    if not data_p3:
        logging.error("❌ Phase 3 thất bại — không có dữ liệu.")
        return {}

    logging.info(f"🔹 Topic (VI): {data_p3.get('meta', {}).get('topic_title_vi', 'N/A')}")
    logging.info("✅ Phase 3 hoàn thành — 4 video scripts (Korean Audio) + word_doc_data OK")
    return data_p3


def run_phase_4(data_p1: dict, data_p2: dict, data_p3: dict) -> dict:
    """
    Phase 4 — DEEP DIVE EPISODE (YouTube Long-form Video).
    
    Input:  data_p1 (News + Đề thi), data_p2 (Essay + Analysis), data_p3 (TikTok scripts)
    Output: JSON cấu trúc cho video_5_deep_dive với các section chi tiết.
    
    RULE: Korean Audio - Vietnamese Subtitles
    Thời lượng mục tiêu: 5-10 phút (YouTube format)
    """
    logging.info("🧠 Phase 4: Deep Dive Episode — Kịch bản YouTube dài...")

    analysis_str = json.dumps(data_p2.get('analysis_list', []), ensure_ascii=False)
    essay_text = data_p2.get('essay', '')
    vocab_data = data_p3.get('word_doc_data', {}).get('vocab_list', [])
    grammar_data = data_p3.get('word_doc_data', {}).get('grammar_list', [])
    
    vocab_str = json.dumps(vocab_data, ensure_ascii=False)
    grammar_str = json.dumps(grammar_data, ensure_ascii=False)

    prompt_p4 = f"""
    Bạn là biên tập viên chương trình "DAILY KOREAN" (데일리 코리안) trên YouTube.
    
    ⛔ NGUYÊN TẮC CỐT LÕI (BẮT BUỘC):
    1. Audio sẽ đọc 100% bằng TIẾNG HÀN.
    2. Tiếng Việt CHỈ dùng làm PHỤ ĐỀ hiển thị trên màn hình.
    3. Thời lượng video mục tiêu: 5-10 phút (YouTube format).
    4. Mỗi segment phải có placeholder "duration_sec": 0 (sẽ được tính sau khi tạo audio).
    
    🚨 QUAN TRỌNG - QUY TẮC NGÔN NGỮ:
    - Tất cả các field kết thúc bằng "_ko" (như explanation_ko, example_ko, analysis_ko) 
      PHẢI VIẾT 100% BẰNG TIẾNG HÀN THUẦN TÚY.
    - TUYỆT ĐỐI KHÔNG được trộn tiếng Việt vào các field "_ko".
    - Tiếng Việt CHỈ được viết trong các field kết thúc bằng "_vi".
    
    Ví dụ ĐÚNG:
    - "explanation_ko": "이 단어는 급격한 변화를 나타내는 표현입니다. 사회나 환경이 빠르게 변할 때 사용합니다."
    - "meaning_vi": "Diễn tả sự thay đổi nhanh chóng, đột ngột"
    
    Ví dụ SAI (❌ KHÔNG LÀM):
    - "explanation_ko": "급변하다. Diễn tả sự thay đổi nhanh chóng..."  ← Có tiếng Việt!
    
    --- DỮ LIỆU ĐẦU VÀO ---
    
    1. [TIN TỨC GỐC]: {data_p1.get('news_summary_easy_kr', '')}
    
    2. [ĐỀ THI TOPIK 54]: {data_p1.get('question_full_text', '')}
    
    3. [VĂN MẪU CHUẨN]: {essay_text}
    
    4. [PHÂN TÍCH CỦA GIÁO SƯ]: {analysis_str}
    
    5. [DANH SÁCH TỪ VỰNG]: {vocab_str}
    
    6. [DANH SÁCH NGỮ PHÁP]: {grammar_str}
    
    --- YÊU CẦU OUTPUT ---
    
    Tạo kịch bản DEEP DIVE cho YouTube với 7 section chi tiết.
    Mỗi section phải có:
    - "ko": Text tiếng Hàn để TTS đọc (100% TIẾNG HÀN, KHÔNG trộn tiếng Việt)
    - "vi": Phụ đề tiếng Việt tương ứng
    - "duration_sec": 0 (placeholder)
    
    OUTPUT JSON (STRICT FORMAT):
    {{
        "video_5_deep_dive": {{
            "meta": {{
                "title_ko": "토픽 54번 완벽 분석 - 오늘의 한국 사회 이슈 (TIẾNG HÀN)",
                "title_vi": "Phân tích đề TOPIK 54 - Vấn đề xã hội Hàn Quốc hôm nay",
                "description_vi": "Mô tả ngắn cho YouTube (~100 từ tiếng Việt)",
                "hashtags": ["#TOPIK", "#KoreanLearning", "#토픽쓰기", "...thêm 5-7 hashtag"]
            }},
            
            "opening": {{
                "hook_ko": "여러분, 토픽 쓰기 54번 문제, 어떻게 준비하고 계세요? (TIẾNG HÀN ~10 giây)",
                "hook_vi": "Các bạn ơi, đề TOPIK 54, các bạn chuẩn bị thế nào rồi?",
                "intro_ko": "안녕하세요, 오늘은 최신 사회 이슈와 토픽 쓰기 54번을 함께 분석해 보겠습니다. (TIẾNG HÀN ~20 giây)",
                "intro_vi": "Xin chào, hôm nay chúng ta sẽ cùng phân tích tin tức xã hội và đề TOPIK 54.",
                "duration_sec": 0
            }},
            
            "news": {{
                "transition_ko": "먼저 오늘의 한국 사회 이슈를 살펴보겠습니다. (TIẾNG HÀN)",
                "transition_vi": "Đầu tiên chúng ta cùng xem tin tức xã hội hôm nay.",
                "content_ko": "Đọc/viết lại tin tức bằng TIẾNG HÀN đơn giản (~30-45 giây, lấy từ input)",
                "content_vi": "Dịch tin tức sang tiếng Việt",
                "analysis_ko": "이 뉴스에서 중요한 점은... (Phân tích bằng TIẾNG HÀN ~30 giây)",
                "analysis_vi": "Điểm quan trọng trong tin này là...",
                "duration_sec": 0
            }},
            
            "transition": {{
                "bridge_ko": "이 주제가 바로 토픽 54번과 연결됩니다. (TIẾNG HÀN ~15 giây)",
                "bridge_vi": "Chủ đề này liên quan trực tiếp đến đề TOPIK 54.",
                "duration_sec": 0
            }},
            
            "exam": {{
                "intro_ko": "이제 토픽 54번 문제를 살펴보겠습니다. (TIẾNG HÀN)",
                "intro_vi": "Bây giờ chúng ta cùng xem đề TOPIK 54.",
                "question_ko": "Lấy đề thi từ input và viết bằng TIẾNG HÀN (~45-60 giây)",
                "question_vi": "Dịch đề thi đầy đủ sang tiếng Việt",
                "tips_ko": "토픽 쓰기에서 중요한 점은 첫째... 둘째... 셋째... (3 tips bằng TIẾNG HÀN)",
                "tips_vi": "Điểm quan trọng khi viết TOPIK là thứ nhất... thứ hai... thứ ba...",
                "duration_sec": 0
            }},
            
            "essay": {{
                "intro_ko": "이제 모범 답안을 함께 읽어보겠습니다. (TIẾNG HÀN - giới thiệu bài văn mẫu)",
                "intro_vi": "Bây giờ chúng ta cùng đọc bài văn mẫu.",
                "paragraphs": [
                    {{
                        "label": "서론 (Mở bài)",
                        "ko": "Nội dung mở bài bằng TIẾNG HÀN (lấy từ VĂN MẪU)",
                        "vi": "Dịch sang tiếng Việt",
                        "analysis_ko": "이 서론에서는... (Phân tích kỹ thuật viết bằng TIẾNG HÀN)",
                        "analysis_vi": "Trong phần mở bài này..."
                    }},
                    {{
                        "label": "본론 1 (Thân bài 1)",
                        "ko": "Nội dung thân bài 1 bằng TIẾNG HÀN",
                        "vi": "Dịch",
                        "analysis_ko": "첫 번째 본론에서는... (TIẾNG HÀN)",
                        "analysis_vi": "Trong thân bài 1..."
                    }},
                    {{
                        "label": "본론 2 (Thân bài 2)",
                        "ko": "Nội dung thân bài 2 bằng TIẾNG HÀN",
                        "vi": "Dịch",
                        "analysis_ko": "두 번째 본론에서는... (TIẾNG HÀN)",
                        "analysis_vi": "Trong thân bài 2..."
                    }},
                    {{
                        "label": "결론 (Kết bài)",
                        "ko": "Nội dung kết bài bằng TIẾNG HÀN",
                        "vi": "Dịch",
                        "analysis_ko": "결론에서는... (TIẾNG HÀN)",
                        "analysis_vi": "Trong phần kết..."
                    }}
                ],
                "duration_sec": 0
            }},
            
            "vocab": {{
                "intro_ko": "이제 오늘 배운 중요한 어휘를 살펴보겠습니다. (TIẾNG HÀN thuần túy)",
                "intro_vi": "Bây giờ chúng ta cùng xem qua từ vựng quan trọng nhé.",
                "items": [
                    {{
                        "word": "급변하다",
                        "explanation_ko": "이 단어는 급격한 변화를 나타내는 표현입니다. 사회나 환경이 빠르게 변할 때 사용합니다. (100% TIẾNG HÀN)",
                        "meaning_vi": "Thay đổi nhanh chóng, biến đổi đột ngột",
                        "example_ko": "세계 경제가 급변하고 있습니다. (100% TIẾNG HÀN)",
                        "example_vi": "Kinh tế thế giới đang thay đổi nhanh chóng."
                    }}
                ],
                "grammar_items": [
                    {{
                        "point": "-로 인해",
                        "explanation_ko": "이 표현은 원인이나 이유를 나타냅니다. 어떤 일의 원인을 설명할 때 사용합니다. (100% TIẾNG HÀN)",
                        "meaning_vi": "Do, vì (nguyên nhân)",
                        "example_ko": "코로나로 인해 많은 변화가 있었습니다. (100% TIẾNG HÀN)",
                        "example_vi": "Do corona, đã có nhiều thay đổi."
                    }}
                ],
                "duration_sec": 0
            }},
            
            "closing": {{
                "summary_ko": "오늘은 토픽 54번 문제와 관련된 내용을 함께 공부했습니다. (TIẾNG HÀN - tóm tắt ~20 giây)",
                "summary_vi": "Hôm nay chúng ta đã cùng học về đề TOPIK 54.",
                "cta_ko": "영상이 도움이 되셨다면 좋아요와 구독 부탁드립니다. 궁금한 점은 댓글로 남겨주세요! (TIẾNG HÀN)",
                "cta_vi": "Nếu video hữu ích, hãy like và subscribe nhé. Có thắc mắc gì để lại comment!",
                "outro_ko": "다음 영상에서 또 만나요! 안녕히 계세요! (TIẾNG HÀN)",
                "outro_vi": "Hẹn gặp lại ở video tiếp theo! Tạm biệt!",
                "duration_sec": 0
            }}
        }}
    }}
    """

    data_p4 = call_ai_api(prompt_p4, temperature=0.7)
    if not data_p4:
        logging.error("❌ Phase 4 thất bại — không có dữ liệu.")
        return {}

    logging.info("✅ Phase 4 hoàn thành — Deep Dive Episode script OK")
    return data_p4


# ==============================================================================
# 3. TIKTOK AUDIO ASSET GENERATION  —  Segment-based MP3 generation
# ==============================================================================

# Cấu hình giọng đọc & tốc độ cho từng video
# RULE: Chỉ sử dụng giọng tiếng Hàn
# Updated to use Azure voice names from AZURE_VOICE_CONFIG
_VOICE_CFG = {
    "video_1": {"voice": "ko-KR-SunHiNeural",   "rate": "-10%",  "role": "news"},      # Healing — nữ, chậm
    "video_2": {"voice": "ko-KR-InJoonNeural",  "rate": "+0%",   "role": "teaching"},  # Teaching — nam
    "video_3": {"voice": "ko-KR-InJoonNeural",  "rate": "+0%",   "role": "exam"},      # Quiz — nam
    "video_4": {"voice": "ko-KR-InJoonNeural",  "rate": "+0%",   "role": "exam"},      # Quiz — nam
    "video_5": {"voice": "ko-KR-JiMinNeural",   "rate": "+0%",   "role": "analysis"},  # Deep Dive — nữ
}

# Thời gian im lặng cho phần "suy nghĩ" trong Quiz (milliseconds)
QUIZ_SILENCE_MS = 4000   # 4 giây


def get_audio_duration(file_path: str) -> float:
    """
    Get audio duration in seconds using mutagen (accurate) or pydub (fallback).
    
    Args:
        file_path: Path to the MP3 file
        
    Returns:
        Duration in seconds (float)
    """
    if not os.path.exists(file_path):
        logging.warning(f"⚠️ Audio file not found: {file_path}")
        return 0.0
    
    try:
        if MUTAGEN_AVAILABLE:
            audio = MP3(file_path)
            return audio.info.length
        else:
            # Fallback to pydub
            audio = AudioSegment.from_file(file_path, format="mp3")
            return len(audio) / 1000.0
    except Exception as e:
        logging.error(f"❌ Error getting audio duration for {file_path}: {e}")
        return 0.0


async def _tts_to_file(text: str, voice: str, rate: str, output_path: str) -> float:
    """
    Generate TTS audio and save directly to file, return duration.
    
    Args:
        text: Korean text to synthesize
        voice: Edge TTS voice name
        rate: Speed rate (e.g., "-10%", "+0%")
        output_path: Path to save the MP3 file
        
    Returns:
        Duration in seconds (float)
    """
    if not text or not text.strip():
        return 0.0
    
    try:
        communicate = edge_tts.Communicate(text.strip(), voice, rate=rate)
        await communicate.save(output_path)
        duration = get_audio_duration(output_path)
        return duration
    except Exception as e:
        logging.error(f"❌ TTS generation failed: {e}")
        return 0.0


async def _tts_to_segment(text: str, voice: str, rate: str) -> AudioSegment:
    """
    Async helper: Gọi edge_tts → lưu tạm → đọc về AudioSegment → xóa file tạm.
    RULE: text PHẢI là tiếng Hàn.
    """
    if not text or not text.strip():
        return AudioSegment.empty()

    tmp_path = os.path.join(
        TEMP_DIR,
        f"_tts_{int(time.time() * 1000)}_{random.randint(1, 99999)}.mp3"
    )

    communicate = edge_tts.Communicate(text.strip(), voice, rate=rate)
    await communicate.save(tmp_path)

    segment = AudioSegment.from_file(tmp_path, format="mp3")

    # Cleanup
    if os.path.exists(tmp_path):
        os.remove(tmp_path)

    return segment


async def _build_video1_news(script: dict, assets_dir: str) -> dict:
    """
    Video 1 — News Healing.
    
    NEW STRUCTURE: opening_ment → content → closing_ment
    Generate separate audio files for each phase with precise timing.
    Auto-compress with SSML +15% if total estimated duration > 55s.
    
    Returns:
        dict with structure:
        {
            "opening": {"audio_path": "/assets/...", "duration": 3.5, "text": "..."},
            "segments": [
                {"ko": "...", "vi": "...", "audio_path": "/assets/...", "duration": 2.5},
                ...
            ],
            "closing": {"audio_path": "/assets/...", "duration": 2.0, "text": "..."},
            "total_duration": 12.5,
            "combined_audio": "/assets/v1_news.mp3"
        }
    """
    cfg = _VOICE_CFG["video_1"]
    
    # Extract script parts
    opening_ment = script.get("opening_ment", "")
    closing_ment = script.get("closing_ment", "다음 영상에서 또 만나요!")
    segments = script.get("segments", [])
    audio_text = script.get("audio_text", "")
    
    # Fallback: if no segments, use audio_text as single segment
    if not segments and audio_text:
        segments = [{"ko": audio_text, "vi": ""}]
    
    if not segments and not opening_ment:
        logging.warning("⚠️  Video 1: No content found — skipping.")
        return {"segments": [], "total_duration": 0, "combined_audio": None}
    
    # Calculate total text for auto-compress check
    total_text = opening_ment + " " + " ".join(s.get("ko", "") for s in segments) + " " + closing_ment
    should_compress, compress_rate = should_compress_audio(total_text)
    base_rate = compress_rate if should_compress else cfg["rate"]
    
    result = {
        "opening": None,
        "segments": [],
        "closing": None,
        "total_duration": 0.0,
        "combined_audio": None,
        "ssml_compressed": should_compress
    }
    
    combined_audio = AudioSegment.empty()
    total_duration = 0.0
    
    # ═══════════════════════════════════════════════════════════════════════════
    # PART 1: Opening Ment
    # ═══════════════════════════════════════════════════════════════════════════
    if opening_ment:
        opening_path = os.path.join(assets_dir, "v1_opening.mp3")
        duration = generate_azure_tts(opening_ment, cfg["voice"], opening_path, base_rate, use_dynamic_rate=False)
        
        if duration > 0:
            result["opening"] = {
                "audio_path": "/assets/v1_opening.mp3",
                "duration": round(duration, 3),
                "text": opening_ment
            }
            total_duration += duration
            
            if os.path.exists(opening_path):
                combined_audio += AudioSegment.from_file(opening_path, format="mp3")
                combined_audio += AudioSegment.silent(duration=300)  # 0.3s pause
                total_duration += 0.3
            
            logging.info(f"🎵 V1 Opening: v1_opening.mp3 ({duration:.2f}s)")
    
    # ═══════════════════════════════════════════════════════════════════════════
    # PART 2: Content Segments
    # ═══════════════════════════════════════════════════════════════════════════
    for idx, seg in enumerate(segments):
        ko_text = seg.get("ko", "")
        vi_text = seg.get("vi", "")
        
        if not ko_text:
            continue
        
        seg_filename = f"v1_seg_{idx}.mp3"
        seg_path = os.path.join(assets_dir, seg_filename)
        
        duration = generate_azure_tts(ko_text, cfg["voice"], seg_path, base_rate, use_dynamic_rate=True)
        
        if duration > 0:
            result["segments"].append({
                "ko": ko_text,
                "vi": vi_text,
                "audio_path": f"/assets/{seg_filename}",
                "duration": round(duration, 3)
            })
            total_duration += duration
            
            if os.path.exists(seg_path):
                combined_audio += AudioSegment.from_file(seg_path, format="mp3")
            
            logging.info(f"🎵 V1 Segment {idx}: {seg_filename} ({duration:.2f}s)")
    
    # ═══════════════════════════════════════════════════════════════════════════
    # PART 3: Closing Ment
    # ═══════════════════════════════════════════════════════════════════════════
    if closing_ment:
        closing_path = os.path.join(assets_dir, "v1_closing.mp3")
        duration = generate_azure_tts(closing_ment, cfg["voice"], closing_path, base_rate, use_dynamic_rate=False)
        
        if duration > 0:
            result["closing"] = {
                "audio_path": "/assets/v1_closing.mp3",
                "duration": round(duration, 3),
                "text": closing_ment
            }
            
            if os.path.exists(closing_path):
                combined_audio += AudioSegment.silent(duration=300)  # 0.3s pause before closing
                combined_audio += AudioSegment.from_file(closing_path, format="mp3")
                total_duration += duration + 0.3
            
            logging.info(f"🎵 V1 Closing: v1_closing.mp3 ({duration:.2f}s)")
    
    # ═══════════════════════════════════════════════════════════════════════════
    # COMBINED AUDIO (backward compatibility)
    # ═══════════════════════════════════════════════════════════════════════════
    combined_path = os.path.join(assets_dir, "v1_news.mp3")
    if len(combined_audio) > 0:
        combined_audio.export(combined_path, format="mp3")
        # Re-measure actual duration
        actual_duration = get_audio_duration(combined_path)
        result["total_duration"] = round(actual_duration, 3)
        result["combined_audio"] = "/assets/v1_news.mp3"
        logging.info(f"🎵 Video 1 combined: v1_news.mp3 ({actual_duration:.1f}s total)")
    else:
        result["total_duration"] = round(total_duration, 3)
    
    return result


async def _build_video2_outline(script: dict, assets_dir: str) -> dict:
    """
    Video 2 — Writing Coach.
    
    NEW STRUCTURE: opening_ment → parts → closing_ment
    Generate separate audio files for each phase with precise timing.
    Auto-compress with SSML +15% if total estimated duration > 55s.
    
    Returns:
        dict with structure:
        {
            "opening": {"audio_path": "/assets/...", "duration": 3.2, "text": "..."},
            "parts": [
                {"role": "intro", "ko": "...", "vi": "...", "audio_path": "/assets/...", "duration": 3.2},
                {"role": "body_1", ...},
                {"role": "body_2", ...},
                {"role": "body_3", ...}
            ],
            "closing": {"audio_path": "/assets/...", "duration": 2.0, "text": "..."},
            "total_duration": 15.5,
            "combined_audio": "/assets/v2_outline.mp3"
        }
    """
    cfg = _VOICE_CFG["video_2"]
    
    # Extract script parts
    opening_ment = script.get("opening_ment", "")
    closing_ment = script.get("closing_ment", "다음 영상에서 또 만나요!")
    parts = script.get("parts", [])
    
    # If no parts, try to build from legacy format
    if not parts:
        legacy_roles = ["intro", "body_1", "body_2", "conclusion"]
        for role in legacy_roles:
            text = script.get(role, "")
            if text:
                parts.append({"role": role, "ko": text, "vi": ""})
    
    if not parts and not opening_ment:
        logging.warning("⚠️  Video 2: No parts found — skipping.")
        return {"parts": [], "total_duration": 0, "combined_audio": None}
    
    # Calculate total text for auto-compress check
    total_text = opening_ment + " " + " ".join(p.get("ko", "") for p in parts) + " " + closing_ment
    should_compress, compress_rate = should_compress_audio(total_text)
    base_rate = compress_rate if should_compress else cfg["rate"]
    
    result = {
        "opening": None,
        "parts": [],
        "closing": None,
        "total_duration": 0.0,
        "combined_audio": None,
        "ssml_compressed": should_compress
    }
    
    combined_audio = AudioSegment.empty()
    total_duration = 0.0
    pause = AudioSegment.silent(duration=500)  # 0.5s between parts
    
    # ═══════════════════════════════════════════════════════════════════════════
    # PART 1: Opening Ment
    # ═══════════════════════════════════════════════════════════════════════════
    if opening_ment:
        opening_path = os.path.join(assets_dir, "v2_opening.mp3")
        duration = generate_azure_tts(opening_ment, cfg["voice"], opening_path, base_rate, use_dynamic_rate=False)
        
        if duration > 0:
            result["opening"] = {
                "audio_path": "/assets/v2_opening.mp3",
                "duration": round(duration, 3),
                "text": opening_ment
            }
            total_duration += duration
            
            if os.path.exists(opening_path):
                combined_audio += AudioSegment.from_file(opening_path, format="mp3")
                combined_audio += pause
                total_duration += 0.5
            
            logging.info(f"🎵 V2 Opening: v2_opening.mp3 ({duration:.2f}s)")
    
    # ═══════════════════════════════════════════════════════════════════════════
    # PART 2: Content Parts (Intro, Body1, Body2, Body3/Conclusion)
    # ═══════════════════════════════════════════════════════════════════════════
    for idx, part in enumerate(parts):
        role = part.get("role", f"part_{idx}")
        ko_text = part.get("ko", "")
        vi_text = part.get("vi", "")
        label_vi = part.get("label_vi", "")
        
        if not ko_text:
            continue
        
        part_filename = f"v2_{role}.mp3"
        part_path = os.path.join(assets_dir, part_filename)
        
        duration = generate_azure_tts(ko_text, cfg["voice"], part_path, base_rate, use_dynamic_rate=True)
        
        if duration > 0:
            result["parts"].append({
                "role": role,
                "label_vi": label_vi,
                "ko": ko_text,
                "vi": vi_text,
                "audio_path": f"/assets/{part_filename}",
                "duration": round(duration, 3)
            })
            total_duration += duration
            
            if os.path.exists(part_path):
                combined_audio += AudioSegment.from_file(part_path, format="mp3")
                if idx < len(parts) - 1:
                    combined_audio += pause
                    total_duration += 0.5
            
            logging.info(f"🎵 V2 {role}: {part_filename} ({duration:.2f}s)")
    
    # ═══════════════════════════════════════════════════════════════════════════
    # PART 3: Closing Ment
    # ═══════════════════════════════════════════════════════════════════════════
    if closing_ment:
        closing_path = os.path.join(assets_dir, "v2_closing.mp3")
        duration = generate_azure_tts(closing_ment, cfg["voice"], closing_path, base_rate, use_dynamic_rate=False)
        
        if duration > 0:
            result["closing"] = {
                "audio_path": "/assets/v2_closing.mp3",
                "duration": round(duration, 3),
                "text": closing_ment
            }
            
            if os.path.exists(closing_path):
                combined_audio += pause
                combined_audio += AudioSegment.from_file(closing_path, format="mp3")
                total_duration += duration + 0.5
            
            logging.info(f"🎵 V2 Closing: v2_closing.mp3 ({duration:.2f}s)")
    
    # ═══════════════════════════════════════════════════════════════════════════
    # COMBINED AUDIO (backward compatibility)
    # ═══════════════════════════════════════════════════════════════════════════
    combined_path = os.path.join(assets_dir, "v2_outline.mp3")
    if len(combined_audio) > 0:
        combined_audio.export(combined_path, format="mp3")
        actual_duration = get_audio_duration(combined_path)
        result["total_duration"] = round(actual_duration, 3)
        result["combined_audio"] = "/assets/v2_outline.mp3"
        logging.info(f"🎵 Video 2 combined: v2_outline.mp3 ({actual_duration:.1f}s total)")
    else:
        result["total_duration"] = round(total_duration, 3)
    
    return result


async def _build_quiz_audio(script: dict, assets_dir: str, video_key: str) -> dict:
    """
    Video 3 & 4 — Quiz (Vocab / Grammar).
    
    NEW STRUCTURE: opening_ment → question → [silence 4s] → answer → closing_ment
    Split audio into separate files for precise Remotion control.
    Auto-compress with SSML +15% if total estimated duration > 55s.
    
    Returns:
        dict with structure:
        {
            "opening_audio": {"path": "/assets/...", "duration": 3.0, "text": "..."},
            "question_audio": {"path": "/assets/...", "duration": 5.2},
            "answer_audio": {"path": "/assets/...", "duration": 8.1},
            "closing_audio": {"path": "/assets/...", "duration": 2.0, "text": "..."},
            "silence_duration": 4.0,
            "total_duration": 22.3,
            "combined_audio": "/assets/v3_vocab_quiz.mp3"
        }
    """
    cfg = _VOICE_CFG[video_key]
    video_num = video_key.split("_")[1]  # "3" or "4"
    
    # Extract data
    opening_ment = script.get("opening_ment", "")
    closing_ment = script.get("closing_ment", "다음 퀴즈에서 또 만나요!")
    question_ko = script.get("question_ko") or script.get("question", "")
    options_ko = script.get("options_ko") or script.get("options", [])
    correct = script.get("correct_answer", "")
    explanation_ko = script.get("explanation_ko") or script.get("explanation", "")
    
    if not question_ko:
        logging.warning(f"⚠️  {video_key}: question_ko rỗng — skipping.")
        return {
            "opening_audio": None,
            "question_audio": None,
            "answer_audio": None,
            "closing_audio": None,
            "silence_duration": QUIZ_SILENCE_MS / 1000.0,
            "total_duration": 0,
            "combined_audio": None
        }
    
    # Calculate total text for auto-compress check
    total_text = (opening_ment + " " + question_ko + " " + " ".join(options_ko) + 
                  " " + f"정답은 {correct}입니다. " + explanation_ko + " " + closing_ment)
    should_compress, compress_rate = should_compress_audio(total_text)
    base_rate = compress_rate if should_compress else cfg["rate"]
    
    result = {
        "opening_audio": None,
        "question_audio": None,
        "answer_audio": None,
        "closing_audio": None,
        "silence_duration": QUIZ_SILENCE_MS / 1000.0,
        "total_duration": 0,
        "combined_audio": None,
        "ssml_compressed": should_compress
    }
    
    combined_audio = AudioSegment.empty()
    total_duration = 0.0
    short_pause = AudioSegment.silent(duration=300)
    
    # ═══════════════════════════════════════════════════════════════════════════
    # PART 0: Opening Ment
    # ═══════════════════════════════════════════════════════════════════════════
    if opening_ment:
        opening_filename = f"v{video_num}_opening.mp3"
        opening_path = os.path.join(assets_dir, opening_filename)
        duration = generate_azure_tts(opening_ment, cfg["voice"], opening_path, base_rate, use_dynamic_rate=False)
        
        if duration > 0:
            result["opening_audio"] = {
                "path": f"/assets/{opening_filename}",
                "duration": round(duration, 3),
                "text": opening_ment
            }
            total_duration += duration
            
            if os.path.exists(opening_path):
                combined_audio += AudioSegment.from_file(opening_path, format="mp3")
                combined_audio += short_pause
                total_duration += 0.3
            
            logging.info(f"🎵 {video_key} opening: {opening_filename} ({duration:.2f}s)")
    
    # ═══════════════════════════════════════════════════════════════════════════
    # PART 1: Question Audio (Question + Options)
    # ═══════════════════════════════════════════════════════════════════════════
    q_filename = f"v{video_num}_question.mp3"
    q_path = os.path.join(assets_dir, q_filename)
    
    # Build question audio: Question + short pause + Options
    q_audio = await _tts_to_segment(question_ko, cfg["voice"], base_rate)
    
    options_text = "  ".join(options_ko)
    opt_audio = await _tts_to_segment(options_text, cfg["voice"], base_rate)
    
    question_combined = q_audio + short_pause + opt_audio
    question_combined.export(q_path, format="mp3")
    q_duration = get_audio_duration(q_path)
    
    result["question_audio"] = {
        "path": f"/assets/{q_filename}",
        "duration": round(q_duration, 3)
    }
    total_duration += q_duration
    combined_audio += question_combined
    
    logging.info(f"🎵 {video_key} question: {q_filename} ({q_duration:.2f}s)")
    
    # ═══════════════════════════════════════════════════════════════════════════
    # SILENCE (4 seconds) - Added to combined, Remotion handles separately
    # ═══════════════════════════════════════════════════════════════════════════
    silence = AudioSegment.silent(duration=QUIZ_SILENCE_MS)
    combined_audio += silence
    total_duration += QUIZ_SILENCE_MS / 1000.0
    
    # ═══════════════════════════════════════════════════════════════════════════
    # PART 2: Answer Audio (Answer announcement + Explanation)
    # ═══════════════════════════════════════════════════════════════════════════
    a_filename = f"v{video_num}_answer.mp3"
    a_path = os.path.join(assets_dir, a_filename)
    
    answer_announce = f"정답은 {correct}입니다."
    ans_audio = await _tts_to_segment(answer_announce, cfg["voice"], base_rate)
    expl_audio = await _tts_to_segment(explanation_ko, cfg["voice"], base_rate)
    
    answer_combined = ans_audio + short_pause + expl_audio
    answer_combined.export(a_path, format="mp3")
    a_duration = get_audio_duration(a_path)
    
    result["answer_audio"] = {
        "path": f"/assets/{a_filename}",
        "duration": round(a_duration, 3)
    }
    total_duration += a_duration
    combined_audio += answer_combined
    
    logging.info(f"🎵 {video_key} answer: {a_filename} ({a_duration:.2f}s)")
    
    # ═══════════════════════════════════════════════════════════════════════════
    # PART 3: Closing Ment
    # ═══════════════════════════════════════════════════════════════════════════
    if closing_ment:
        closing_filename = f"v{video_num}_closing.mp3"
        closing_path = os.path.join(assets_dir, closing_filename)
        duration = generate_azure_tts(closing_ment, cfg["voice"], closing_path, base_rate, use_dynamic_rate=False)
        
        if duration > 0:
            result["closing_audio"] = {
                "path": f"/assets/{closing_filename}",
                "duration": round(duration, 3),
                "text": closing_ment
            }
            
            if os.path.exists(closing_path):
                combined_audio += short_pause
                combined_audio += AudioSegment.from_file(closing_path, format="mp3")
                total_duration += duration + 0.3
            
            logging.info(f"🎵 {video_key} closing: {closing_filename} ({duration:.2f}s)")
    
    # ═══════════════════════════════════════════════════════════════════════════
    # COMBINED AUDIO (backward compatibility)
    # ═══════════════════════════════════════════════════════════════════════════
    combined_filename = f"v{video_num}_{'vocab' if video_num == '3' else 'grammar'}_quiz.mp3"
    combined_path = os.path.join(assets_dir, combined_filename)
    combined_audio.export(combined_path, format="mp3")
    
    actual_duration = get_audio_duration(combined_path)
    result["total_duration"] = round(actual_duration, 3)
    result["combined_audio"] = f"/assets/{combined_filename}"
    
    logging.info(f"🎵 {video_key} combined: {combined_filename} ({actual_duration:.1f}s total)")
    
    return result


async def _build_video5_deep_dive(script: dict, assets_dir: str) -> dict:
    """
    Video 5 — Deep Dive Episode (YouTube Long-form).
    
    Generate separate MP3 for EACH segment with precise timing.
    Uses Azure TTS with JiMinNeural voice for analysis style.
    
    Returns:
        dict with structure:
        {
            "segments": [
                {"section": "opening", "ko": "...", "vi": "...", "audio_path": "/assets/deep_0.mp3", "duration": 15.2},
                {"section": "news", ...},
                ...
            ],
            "total_duration": 420.5,  # ~7 minutes
            "combined_audio": "/assets/v5_deep_dive.mp3",
            "timestamps": [
                {"section": "opening", "start_sec": 0, "label": "🎬 Intro"},
                {"section": "news", "start_sec": 32, "label": "📰 Tin tức"},
                ...
            ]
        }
    """
    cfg = _VOICE_CFG["video_5"]
    
    if not script:
        logging.warning("⚠️ Video 5: No deep_dive script found — skipping.")
        return {"segments": [], "total_duration": 0, "combined_audio": None, "timestamps": []}
    
    result_segments = []
    combined_audio = AudioSegment.empty()
    timestamps = []
    total_duration = 0.0
    segment_idx = 0
    pause = AudioSegment.silent(duration=500)  # 0.5s between sections
    
    # Voice assignments for different parts of Deep Dive
    voice_host = AZURE_VOICE_CONFIG.get("host", "ko-KR-SunHiNeural")
    voice_news = AZURE_VOICE_CONFIG.get("news", "ko-KR-SunHiNeural")
    voice_exam = AZURE_VOICE_CONFIG.get("exam", "ko-KR-InJoonNeural")
    voice_analysis = AZURE_VOICE_CONFIG.get("analysis", "ko-KR-JiMinNeural")
    
    async def process_segment(section_name: str, ko_text: str, vi_text: str, voice: str, rate: str = "+0%"):
        """Helper to process a single segment."""
        nonlocal segment_idx, total_duration, combined_audio
        
        if not ko_text or not ko_text.strip():
            return
        
        # Generate individual audio file
        seg_filename = f"deep_{segment_idx}.mp3"
        seg_path = os.path.join(assets_dir, seg_filename)
        
        # Use Azure TTS (or fallback to edge-tts)
        duration = await generate_azure_tts_async(ko_text, voice, seg_path, rate)
        
        if duration <= 0:
            logging.warning(f"⚠️ Deep Dive segment {segment_idx} ({section_name}): TTS failed, skipping.")
            return
        
        relative_path = f"/assets/{seg_filename}"
        
        result_segments.append({
            "section": section_name,
            "ko": ko_text,
            "vi": vi_text,
            "audio_path": relative_path,
            "duration": round(duration, 3)
        })
        
        # Add timestamp marker
        timestamps.append({
            "section": section_name,
            "start_sec": round(total_duration, 0),
            "label": _get_section_label(section_name)
        })
        
        total_duration += duration
        
        # Build combined audio
        if os.path.exists(seg_path):
            seg_audio = AudioSegment.from_file(seg_path, format="mp3")
            combined_audio += seg_audio + pause
            total_duration += 0.5  # Account for pause
        
        logging.info(f"🎵 Deep Dive [{section_name}]: {seg_filename} ({duration:.2f}s)")
        segment_idx += 1
    
    # ═══════════════════════════════════════════════════════════════════════════
    # Process each section of the Deep Dive script
    # ═══════════════════════════════════════════════════════════════════════════
    
    # 1. OPENING
    opening = script.get("opening", {})
    if opening:
        hook_ko = opening.get("hook_ko", "")
        intro_ko = opening.get("intro_ko", "")
        combined_opening = f"{hook_ko} {intro_ko}".strip()
        combined_vi = f"{opening.get('hook_vi', '')} {opening.get('intro_vi', '')}".strip()
        await process_segment("opening", combined_opening, combined_vi, voice_host, "-5%")
    
    # 2. NEWS
    news = script.get("news", {})
    if news:
        news_parts = []
        vi_parts = []
        for key in ["transition_ko", "content_ko", "analysis_ko"]:
            if news.get(key):
                news_parts.append(news[key])
        for key in ["transition_vi", "content_vi", "analysis_vi"]:
            if news.get(key):
                vi_parts.append(news[key])
        combined_news_ko = " ".join(news_parts)
        combined_news_vi = " ".join(vi_parts)
        await process_segment("news", combined_news_ko, combined_news_vi, voice_news)
    
    # 3. TRANSITION
    transition = script.get("transition", {})
    if transition:
        await process_segment(
            "transition",
            transition.get("bridge_ko", ""),
            transition.get("bridge_vi", ""),
            voice_host
        )
    
    # 4. EXAM
    exam = script.get("exam", {})
    if exam:
        exam_parts = []
        vi_parts = []
        for key in ["intro_ko", "question_ko", "tips_ko"]:
            if exam.get(key):
                exam_parts.append(exam[key])
        for key in ["intro_vi", "question_vi", "tips_vi"]:
            if exam.get(key):
                vi_parts.append(exam[key])
        combined_exam_ko = " ".join(exam_parts)
        combined_exam_vi = " ".join(vi_parts)
        await process_segment("exam", combined_exam_ko, combined_exam_vi, voice_exam, "-5%")
    
    # 5. ESSAY (Process each paragraph separately for better timestamps)
    essay = script.get("essay", {})
    if essay:
        # Essay intro
        intro_ko = essay.get("intro_ko", "")
        intro_vi = essay.get("intro_vi", "")
        if intro_ko:
            await process_segment("essay_intro", intro_ko, intro_vi, voice_analysis)
        
        # Essay paragraphs
        paragraphs = essay.get("paragraphs", [])
        for idx, para in enumerate(paragraphs):
            label = para.get("label", f"Para {idx+1}")
            para_ko = para.get("ko", "")
            para_vi = para.get("vi", "")
            analysis_ko = para.get("analysis_ko", "")
            analysis_vi = para.get("analysis_vi", "")
            
            # Combine paragraph content with analysis
            combined_ko = f"{para_ko} {analysis_ko}".strip()
            combined_vi = f"{para_vi} {analysis_vi}".strip()
            
            if combined_ko:
                await process_segment(f"essay_{label}", combined_ko, combined_vi, voice_analysis)
    
    # 6. VOCAB
    vocab = script.get("vocab", {})
    if vocab:
        # Vocab intro
        vocab_intro = vocab.get("intro_ko", "")
        vocab_intro_vi = vocab.get("intro_vi", "")
        if vocab_intro:
            await process_segment("vocab_intro", vocab_intro, vocab_intro_vi, voice_analysis)
        
        # Vocab items
        vocab_items = vocab.get("items", [])
        for item in vocab_items:
            word = item.get("word", "")
            explanation_ko = item.get("explanation_ko", "")
            example_ko = item.get("example_ko", "")
            meaning_vi = item.get("meaning_vi", "")
            example_vi = item.get("example_vi", "")
            
            combined_ko = f"{word}. {explanation_ko} {example_ko}".strip()
            combined_vi = f"{meaning_vi} {example_vi}".strip()
            
            if combined_ko:
                await process_segment(f"vocab_{word}", combined_ko, combined_vi, voice_analysis)
        
        # Grammar items
        grammar_items = vocab.get("grammar_items", [])
        for item in grammar_items:
            point = item.get("point", "")
            explanation_ko = item.get("explanation_ko", "")
            example_ko = item.get("example_ko", "")
            meaning_vi = item.get("meaning_vi", "")
            example_vi = item.get("example_vi", "")
            
            combined_ko = f"{point}. {explanation_ko} {example_ko}".strip()
            combined_vi = f"{meaning_vi} {example_vi}".strip()
            
            if combined_ko:
                await process_segment(f"grammar_{point}", combined_ko, combined_vi, voice_analysis)
    
    # 7. CLOSING
    closing = script.get("closing", {})
    if closing:
        closing_parts = []
        vi_parts = []
        for key in ["summary_ko", "cta_ko", "outro_ko"]:
            if closing.get(key):
                closing_parts.append(closing[key])
        for key in ["summary_vi", "cta_vi", "outro_vi"]:
            if closing.get(key):
                vi_parts.append(closing[key])
        combined_closing_ko = " ".join(closing_parts)
        combined_closing_vi = " ".join(vi_parts)
        await process_segment("closing", combined_closing_ko, combined_closing_vi, voice_host)
    
    # ═══════════════════════════════════════════════════════════════════════════
    # Export combined audio
    # ═══════════════════════════════════════════════════════════════════════════
    combined_path = os.path.join(assets_dir, "v5_deep_dive.mp3")
    if len(combined_audio) > 0:
        combined_audio.export(combined_path, format="mp3")
        logging.info(f"🎵 Video 5 combined: {combined_path} ({total_duration:.1f}s = {total_duration/60:.1f}min total)")
    
    return {
        "segments": result_segments,
        "total_duration": round(total_duration, 3),
        "combined_audio": "/assets/v5_deep_dive.mp3",
        "timestamps": timestamps
    }


def _get_section_label(section: str) -> str:
    """Get human-readable label for timestamp."""
    labels = {
        "opening": "🎬 Intro",
        "news": "📰 Tin tức",
        "transition": "🔄 Chuyển tiếp",
        "exam": "📝 Đề thi TOPIK 54",
        "essay_intro": "✍️ Văn mẫu",
        "vocab_intro": "📚 Từ vựng & Ngữ pháp",
        "closing": "👋 Kết thúc",
    }
    # Check for partial matches
    for key, label in labels.items():
        if section.startswith(key):
            return label
    return section


async def generate_tiktok_assets(phase3_json: dict, assets_dir: str, phase4_json: dict = None) -> dict:
    """
    Entry-point for audio asset generation with SEGMENT-LEVEL timing.

    Input:  phase3_json — data from Phase 3 (contains tiktok_script)
    
    Output: dict with structure:
        {
            "audio_paths": {
                "video_1_news": "/assets/v1_news.mp3",
                "video_2_outline": "/assets/v2_outline.mp3",
                ...
            },
            "audio_data": {
                "video_1_news": {
                    "segments": [{...}],
                    "total_duration": 15.5,
                    ...
                },
                ...
            }
        }
    
    The audio_data contains precise timing info for each segment/part.
    """
    logging.info("🎤 Bắt đầu generate_tiktok_assets — Segment-based audio generation...")

    tiktok = phase3_json.get("tiktok_script", {})
    if not tiktok:
        logging.error("❌ tiktok_script không tìm thấy trong Phase 3 output.")
        return {"audio_paths": {}, "audio_data": {}}

    os.makedirs(assets_dir, exist_ok=True)

    audio_paths = {}   # Backward compatible: {video_key: combined_audio_path}
    audio_data = {}    # NEW: Detailed timing data per video

    # ═══════════════════════════════════════════════════════════════════════════
    # Video 1: News Healing — Per-segment audio
    # ═══════════════════════════════════════════════════════════════════════════
    v1_data = await _build_video1_news(tiktok.get("video_1_news", {}), assets_dir)
    audio_data["video_1_news"] = v1_data
    if v1_data.get("combined_audio"):
        audio_paths["video_1_news"] = os.path.join(assets_dir, "v1_news.mp3")

    # ═══════════════════════════════════════════════════════════════════════════
    # Video 2: Writing Coach — Per-part audio
    # ═══════════════════════════════════════════════════════════════════════════
    v2_data = await _build_video2_outline(tiktok.get("video_2_outline", {}), assets_dir)
    audio_data["video_2_outline"] = v2_data
    if v2_data.get("combined_audio"):
        audio_paths["video_2_outline"] = os.path.join(assets_dir, "v2_outline.mp3")

    # ═══════════════════════════════════════════════════════════════════════════
    # Video 3: Vocab Quiz — Question + Answer split
    # ═══════════════════════════════════════════════════════════════════════════
    v3_data = await _build_quiz_audio(tiktok.get("video_3_vocab_quiz", {}), assets_dir, "video_3")
    audio_data["video_3_vocab_quiz"] = v3_data
    if v3_data.get("combined_audio"):
        audio_paths["video_3_vocab_quiz"] = os.path.join(assets_dir, "v3_vocab_quiz.mp3")

    # ═══════════════════════════════════════════════════════════════════════════
    # Video 4: Grammar Quiz — Question + Answer split
    # ═══════════════════════════════════════════════════════════════════════════
    v4_data = await _build_quiz_audio(tiktok.get("video_4_grammar_quiz", {}), assets_dir, "video_4")
    audio_data["video_4_grammar_quiz"] = v4_data
    if v4_data.get("combined_audio"):
        audio_paths["video_4_grammar_quiz"] = os.path.join(assets_dir, "v4_grammar_quiz.mp3")

    # ═══════════════════════════════════════════════════════════════════════════
    # Video 5: Deep Dive Episode — Per-segment audio (YouTube long-form)
    # ═══════════════════════════════════════════════════════════════════════════
    if phase4_json and phase4_json.get("video_5_deep_dive"):
        v5_data = await _build_video5_deep_dive(phase4_json["video_5_deep_dive"], assets_dir)
        audio_data["video_5_deep_dive"] = v5_data
        if v5_data.get("combined_audio"):
            audio_paths["video_5_deep_dive"] = os.path.join(assets_dir, "v5_deep_dive.mp3")
    else:
        logging.info("ℹ️ Video 5 (Deep Dive) skipped — no Phase 4 data provided.")

    logging.info("✅ generate_tiktok_assets hoàn thành — Segment-based audio với timing chính xác.")
    
    return {
        "audio_paths": audio_paths,
        "audio_data": audio_data
    }


# ==============================================================================
# 4. WORD DOCUMENT CREATION  (Giữ nguyên logic, cập nhật data source)
# ==============================================================================

def create_professional_docx(data_p1: dict, data_p2: dict, data_p3: dict, source_url: str) -> str | None:
    """
    Tạo file Word chuyên nghiệp từ dữ liệu 3 phases.
    Dữ liệu vocab/grammar lấy từ word_doc_data (Phase 3 mới).
    """
    logging.info("📝 Đang tạo file Word...")

    # --- Xử lý tên file ---
    raw_title = data_p1.get('topic_korean', 'Topic_Moi')
    safe_title = re.sub(r'[\\/*?:"<>|]', "", raw_title).replace(" ", "_")

    output_dir = os.environ.get('OUTPUT_DIR', 'public')
    os.makedirs(output_dir, exist_ok=True)
    filename = os.path.join(output_dir, f"TOPIK_WRITING_{safe_title[:30]}.docx")

    try:
        doc = Document()

        # Default font
        style      = doc.styles['Normal']
        style.font.name = 'Times New Roman'
        style.font.size = Pt(11)

        # ===== HEADER =====
        header_text = f"TOPIK II CÂU 54 - {sanitize_text(raw_title)}"
        header = doc.add_heading(header_text, 0)
        header.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph(f"Nguồn tin: {source_url}").italic = True
        doc.add_paragraph(f"Ngày tạo: {datetime.now().strftime('%d/%m/%Y')}")
        doc.add_paragraph("-" * 60)

        # ===== 1. ĐỀ BÀI =====
        doc.add_heading('1. ĐỀ BÀI (QUESTION)', level=1)
        table = doc.add_table(rows=1, cols=1)
        table.style = 'Table Grid'
        cell = table.cell(0, 0)

        korean_question = data_p1.get('question_full_text', '')
        p_kr = cell.add_paragraph(sanitize_text(korean_question))
        p_kr.paragraph_format.space_after = Pt(12)

        doc.add_paragraph().add_run("\n")

        # ===== 2. BÀI VĂN MẪU =====
        doc.add_heading('2. BÀI VĂN MẪU (MODEL ESSAY)', level=1)

        essay_kr = data_p2.get('essay', '')
        doc.add_heading('🇰🇷 Tiếng Hàn:', level=3)

        for para in essay_kr.split('\n'):
            if para.strip():
                p = doc.add_paragraph(sanitize_text(para.strip()))
                p.paragraph_format.first_line_indent = Pt(20)
                p.paragraph_format.space_after = Pt(6)

        doc.add_paragraph().add_run("\n")

        # ===== 3. TỪ VỰNG & NGỮ PHÁP (từ word_doc_data) =====
        doc.add_heading('3. TỪ VỰNG (VOCABULARY)', level=1)

        word_doc = data_p3.get('word_doc_data', {})
        vocab_list = word_doc.get('vocab_list', [])

        if not vocab_list:
            doc.add_paragraph("(Không có dữ liệu từ vựng)")
        else:
            for item in vocab_list:
                word       = item.get('word', 'N/A')
                meaning_vi = item.get('meaning_vi', '')
                example    = item.get('example', '')

                p = doc.add_paragraph(style='List Bullet')

                # Từ vựng — in đậm, xanh
                run_word = p.add_run(sanitize_text(word))
                run_word.bold = True
                run_word.font.color.rgb = RGBColor(0, 50, 150)

                if meaning_vi:
                    p.add_run(f" : {sanitize_text(meaning_vi)}")

                if example:
                    run_ex = p.add_run(f"\n   └ 💡 {sanitize_text(example)}")
                    run_ex.font.size = Pt(10)
                    run_ex.font.color.rgb = RGBColor(100, 100, 100)

        # --- Ngữ pháp ---
        doc.add_paragraph().add_run("\n")
        doc.add_heading('4. NGỮ PHÁP (GRAMMAR)', level=1)

        grammar_list = word_doc.get('grammar_list', [])

        if not grammar_list:
            doc.add_paragraph("(Không có dữ liệu ngữ pháp)")
        else:
            for item in grammar_list:
                point      = item.get('point', 'N/A')
                meaning_vi = item.get('meaning_vi', '')
                example    = item.get('example', '')

                p = doc.add_paragraph(style='List Bullet')

                run_point = p.add_run(sanitize_text(point))
                run_point.bold = True
                run_point.font.color.rgb = RGBColor(0, 100, 50)

                if meaning_vi:
                    p.add_run(f" : {sanitize_text(meaning_vi)}")

                if example:
                    run_ex = p.add_run(f"\n   └ 💡 {sanitize_text(example)}")
                    run_ex.font.size = Pt(10)
                    run_ex.font.color.rgb = RGBColor(100, 100, 100)

        # --- Cloze Test ---
        doc.add_paragraph().add_run("\n")
        doc.add_heading('5. CLOZE TEST (ĐIỀN CHỖ TRỐNG)', level=1)

        cloze = word_doc.get('cloze_test', {})
        if cloze:
            doc.add_paragraph(f"📝 Câu hỏi: {sanitize_text(cloze.get('question', ''))}")
            doc.add_paragraph(f"✅ Đáp án:  {sanitize_text(cloze.get('answer', ''))}")
            doc.add_paragraph(f"💡 Gợi ý:   {sanitize_text(cloze.get('hint_vi', ''))}")

        # ===== LƯU FILE =====
        try:
            doc.save(filename)
            logging.info(f"✅ Đã tạo file Word: {filename}")
            return filename
        except PermissionError:
            new_filename = filename.replace(".docx", f"_{int(time.time())}.docx")
            logging.warning(f"⚠️  File đang mở. Lưu sang {new_filename}...")
            doc.save(new_filename)
            return new_filename

    except Exception as e:
        logging.error(f"❌ Lỗi tạo Word: {e}")
        traceback.print_exc()
        return None


# ==============================================================================
# 5. YOUTUBE METADATA GENERATION  —  Auto-generate timestamps, title, hashtags
# ==============================================================================

def generate_youtube_description(json_data: dict, output_path: str = None) -> str:
    """
    Generate YouTube description with auto-calculated timestamps.
    
    Args:
        json_data: Final data JSON containing audio_data with timestamps
        output_path: Path to save youtube_info.txt (optional)
        
    Returns:
        Generated description text
    """
    logging.info("📝 Đang tạo YouTube metadata...")
    
    # Extract data
    meta = json_data.get("meta", {})
    phase1 = json_data.get("phase1", {})
    phase4 = json_data.get("phase4", {})
    audio_data = json_data.get("audio_data", {})
    
    # Get video 5 timestamps
    v5_data = audio_data.get("video_5_deep_dive", {})
    timestamps = v5_data.get("timestamps", [])
    total_duration = v5_data.get("total_duration", 0)
    
    # Get metadata from Phase 4
    deep_dive_meta = {}
    if phase4 and phase4.get("video_5_deep_dive"):
        deep_dive_meta = phase4["video_5_deep_dive"].get("meta", {})
    
    # ═══════════════════════════════════════════════════════════════════════════
    # BUILD YOUTUBE DESCRIPTION
    # ═══════════════════════════════════════════════════════════════════════════
    
    lines = []
    
    # Title
    title_vi = deep_dive_meta.get("title_vi", meta.get("topic_title_vi", "TOPIK 쓰기 54 - Deep Dive"))
    title_ko = deep_dive_meta.get("title_ko", phase1.get("topic_korean", ""))
    
    lines.append("=" * 60)
    lines.append("📺 YOUTUBE VIDEO METADATA")
    lines.append("=" * 60)
    lines.append("")
    lines.append(f"🎬 TITLE (VI): {title_vi}")
    lines.append(f"🎬 TITLE (KO): {title_ko}")
    lines.append("")
    
    # Duration
    duration_min = total_duration / 60
    lines.append(f"⏱️ DURATION: {_format_timestamp(total_duration)} ({duration_min:.1f} phút)")
    lines.append("")
    
    # Timestamps
    lines.append("📌 TIMESTAMPS:")
    lines.append("-" * 40)
    
    if timestamps:
        for ts in timestamps:
            start_sec = ts.get("start_sec", 0)
            label = ts.get("label", ts.get("section", ""))
            timestamp_str = _format_timestamp(start_sec)
            lines.append(f"{timestamp_str} - {label}")
    else:
        # Fallback: Generate estimated timestamps
        lines.append("00:00 - 🎬 Intro")
        lines.append("00:30 - 📰 Tin tức & Phân tích")
        lines.append("02:00 - 📝 Đề thi TOPIK 54")
        lines.append("03:30 - ✍️ Văn mẫu & Phân tích")
        lines.append("06:00 - 📚 Từ vựng & Ngữ pháp")
        lines.append("08:00 - 👋 Kết thúc")
    
    lines.append("")
    
    # Description
    lines.append("📄 DESCRIPTION:")
    lines.append("-" * 40)
    
    description = deep_dive_meta.get("description_vi", "")
    if description:
        lines.append(description)
    else:
        # Generate default description
        topic = meta.get("topic_title_vi", "")
        lines.append(f"🎓 DAILY KOREAN - 데일리 코리안 | Phân tích chuyên sâu đề thi TOPIK II Câu 54")
        lines.append(f"")
        lines.append(f"Trong video này, chúng ta sẽ cùng nhau:")
        lines.append(f"✅ Phân tích xu hướng tin tức xã hội Hàn Quốc")
        lines.append(f"✅ Làm quen với dạng đề TOPIK 54 (600-700 chữ)")
        lines.append(f"✅ Học cách viết bài văn mẫu đạt điểm tối đa")
        lines.append(f"✅ Nắm vững từ vựng và ngữ pháp quan trọng")
        lines.append(f"")
        lines.append(f"📚 Chủ đề hôm nay: {topic}")
    
    lines.append("")
    
    # Hashtags
    lines.append("🏷️ HASHTAGS:")
    lines.append("-" * 40)
    
    hashtags = deep_dive_meta.get("hashtags", [])
    if not hashtags:
        hashtags = [
            "#TOPIK", "#TOPIKwriting", "#토픽쓰기", "#토픽54",
            "#KoreanLearning", "#LearnKorean", "#HọcTiếngHàn",
            "#TOPIKII", "#KoreanTest", "#토픽시험",
            "#DailyKorean", "#데일리코리안", "#VietnamKorea"
        ]
    
    lines.append(" ".join(hashtags))
    lines.append("")
    
    # SEO Keywords
    lines.append("🔍 SEO KEYWORDS:")
    lines.append("-" * 40)
    lines.append("TOPIK 쓰기 54, TOPIK writing, 토픽 작문, học tiếng Hàn, Korean essay, mẫu bài viết TOPIK")
    lines.append("")
    
    # Social Links (placeholder)
    lines.append("🔗 LINKS:")
    lines.append("-" * 40)
    lines.append("📱 TikTok: @deep_dive_korean")
    lines.append("📸 Instagram: @deep_dive_korean")
    lines.append("💬 Discord: [Link cộng đồng]")
    lines.append("")
    
    lines.append("=" * 60)
    
    # Join all lines
    description_text = "\n".join(lines)
    
    # Save to file if path provided
    if output_path:
        try:
            with open(output_path, "w", encoding="utf-8") as f:
                f.write(description_text)
            logging.info(f"✅ Đã lưu YouTube metadata: {output_path}")
        except Exception as e:
            logging.error(f"❌ Lỗi lưu YouTube metadata: {e}")
    
    return description_text


def _format_timestamp(seconds: float) -> str:
    """Format seconds as MM:SS or HH:MM:SS timestamp."""
    seconds = int(seconds)
    hours = seconds // 3600
    minutes = (seconds % 3600) // 60
    secs = seconds % 60
    
    if hours > 0:
        return f"{hours:02d}:{minutes:02d}:{secs:02d}"
    else:
        return f"{minutes:02d}:{secs:02d}"


# ==============================================================================
# 6. REMOTION RENDER  —  Hỗ trợ CompositionID tùy chọn + Error Handling
# ==============================================================================

def render_single_video(composition_id: str, json_path: str, output_path: str) -> bool:
    """
    Render 1 video với CompositionID cụ thể.
    Sử dụng đường dẫn tuyệt đối cho json_path và output_path.
    """
    abs_json   = os.path.abspath(json_path)
    abs_output = os.path.abspath(output_path)

    if not os.path.exists(abs_json):
        logging.error(f"❌ [{composition_id}] File data không tồn tại: {abs_json}")
        return False

    logging.info(f"🎥 Render [{composition_id}] → {os.path.basename(abs_output)}")

    # Trên Windows, cần sử dụng shell=True để tìm npx trong PATH
    # Trên Linux, shell=False hoạt động tốt hơn (tránh TTY issues)
    import platform
    is_windows = platform.system() == "Windows"
    
    cmd = [
        "npx", "remotion", "render",
        composition_id,             # Tên Composition (khác nhau cho mỗi video)
        abs_output,                 # Output path (tuyệt đối)
        "--props", abs_json,        # Props JSON (tuyệt đối)
        "--concurrency=1",          # Concurrency=1 để tránh browser crash
        "--gl=angle" if is_windows else "--gl=swangle",  # swangle tốt hơn trên Linux headless
        "--log=info"
    ]

    try:
        # shell=True cần thiết trên Windows để tìm npx trong PATH
        # shell=False trên Linux để tránh TTY/interactive shell issues
        subprocess.run(
            cmd, 
            check=True, 
            cwd="topik-video", 
            capture_output=False, 
            shell=is_windows,
            stdin=subprocess.DEVNULL  # Prevent interactive prompts
        )

        if os.path.exists(abs_output):
            file_size_mb = os.path.getsize(abs_output) / (1024 * 1024)
            logging.info(f"✅ [{composition_id}] OK — {file_size_mb:.1f} MB")
            return True
        else:
            logging.error(f"❌ [{composition_id}] Render done nhưng file không thấy.")
            return False

    except subprocess.CalledProcessError as e:
        logging.error(f"❌ [{composition_id}] Remotion lỗi (Exit {e.returncode}). Xem log bên trên.")
        return False
    except Exception as e:
        logging.error(f"❌ [{composition_id}] Lỗi ngoại lệ: {e}")
        return False


def render_all_videos(json_path: str, include_deep_dive: bool = True) -> list[str]:
    """
    Render loop: chạy 5 lần liên tiếp, mỗi lần với CompositionID khác nhau.
    
    Args:
        json_path: Path to the final_data.json file
        include_deep_dive: Whether to include Video 5 (Deep Dive)
        
    Returns: danh sách paths của các video đã render thành công.
    
    Error Handling: Mỗi video được wrap trong try/except riêng để đảm bảo
    nếu 1 video lỗi, các video khác vẫn tiếp tục render.
    """
    logging.info("=" * 60)
    logging.info("🎬 BẮT ĐẦU RENDER LOOP — 5 VIDEO (4 TikTok + 1 YouTube Deep Dive)")
    logging.info("=" * 60)

    timestamp   = int(time.time())
    rendered    = []                    # Paths video thành công
    failed      = []                    # Composition IDs that failed
    
    manifest_to_use = VIDEO_MANIFEST if include_deep_dive else VIDEO_MANIFEST[:4]

    for entry in manifest_to_use:
        composition = entry["composition"]    # e.g. "TikTok_NewsHealing"
        prefix      = entry["prefix"]         # e.g. "V1_News"

        video_filename = f"{prefix}_{timestamp}.mp4"
        video_path     = os.path.join("topik-video", "public", video_filename)

        try:
            success = render_single_video(composition, json_path, video_path)

            if success:
                rendered.append(video_path)
                logging.info(f"✅ [{composition}] Render thành công.")
            else:
                failed.append(composition)
                logging.warning(f"⚠️ [{composition}] thất bại — tiếp tục các video còn lại.")
                
        except Exception as e:
            failed.append(composition)
            logging.error(f"❌ [{composition}] Exception during render: {e}")
            traceback.print_exc()
            logging.info("   → Tiếp tục với video tiếp theo...")
            continue

    # Summary
    total = len(manifest_to_use)
    success_count = len(rendered)
    fail_count = len(failed)
    
    logging.info("=" * 60)
    logging.info(f"🎬 Render loop hoàn thành: {success_count}/{total} video OK")
    if failed:
        logging.warning(f"   ❌ Failed compositions: {', '.join(failed)}")
    logging.info("=" * 60)
    
    return rendered


# ==============================================================================
# 7. MAIN — Orchestrator (Updated for 5 videos + YouTube metadata)
# ==============================================================================

def main():
    logging.info("=" * 60)
    logging.info("🚀 DAILY KOREAN v3.0 — 데일리 코리안 Content Automation")
    logging.info("=" * 60)

    # Biến để upload sau cùng
    docx_path    = None
    youtube_info_path = None
    rendered_videos = []

    # ------------------------------------------------------------------
    # PHASE 1: Crawl News + Ra đề
    # ------------------------------------------------------------------
    url_rss, source_name = get_latest_editorial_rss()
    if not url_rss:
        logging.error("❌ Không tìm được bài báo từ RSS. Dừng.")
        return

    content = extract_content(url_rss)
    if not content:
        logging.error("❌ Không tải được nội dung bài báo. Dừng.")
        return

    data_p1 = run_phase_1(content['text'])
    if not data_p1:
        logging.error("❌ Phase 1 thất bại. Dừng.")
        return

    # ------------------------------------------------------------------
    # PHASE 2: Văn mẫu + Phân tích
    # ------------------------------------------------------------------
    data_p2 = run_phase_2(data_p1)
    if not data_p2:
        logging.error("❌ Phase 2 thất bại. Dừng.")
        return

    # ------------------------------------------------------------------
    # Tải video nền (song song với Phase 3 — không block)
    # ------------------------------------------------------------------
    keyword = data_p1.get('video_keyword', 'study')
    bg_download_result = download_background_video(keyword, os.path.join(ASSETS_DIR, "background_loop.mp4"))
    
    # Extract video duration from download result
    video_bg_duration = 0.0
    if isinstance(bg_download_result, dict):
        video_bg_duration = bg_download_result.get("duration", 0.0)
    elif VIDEO_BG_DURATION_CACHE > 0:
        video_bg_duration = VIDEO_BG_DURATION_CACHE

    # ------------------------------------------------------------------
    # PHASE 3: Multi-channel editor → JSON 4 video + Word data
    # ------------------------------------------------------------------
    data_p3 = run_phase_3(data_p1, data_p2)
    if not data_p3:
        logging.error("❌ Phase 3 thất bại. Dừng.")
        return

    # ------------------------------------------------------------------
    # PHASE 4: Deep Dive Episode → JSON video 5 (YouTube long-form)
    # ------------------------------------------------------------------
    data_p4 = run_phase_4(data_p1, data_p2, data_p3)
    if not data_p4:
        logging.warning("⚠️ Phase 4 thất bại — Video 5 sẽ bị bỏ qua.")
        include_deep_dive = False
    else:
        include_deep_dive = True
        logging.info("✅ Phase 4 hoàn thành — Deep Dive script OK")

    # ------------------------------------------------------------------
    # GENERATE TIKTOK AUDIO ASSETS — Segment-based with timing (5 videos)
    # ------------------------------------------------------------------
    audio_result = asyncio.run(generate_tiktok_assets(data_p3, ASSETS_DIR, data_p4 if include_deep_dive else None))
    if not audio_result or not audio_result.get("audio_paths"):
        logging.error("❌ Không tạo được audio assets. Dừng.")
        return
    
    audio_paths = audio_result["audio_paths"]
    audio_data = audio_result["audio_data"]

    # ------------------------------------------------------------------
    # TẠO FILE WORD
    # ------------------------------------------------------------------
    docx_path = create_professional_docx(data_p1, data_p2, data_p3, url_rss)

    # ------------------------------------------------------------------
    # LƯU final_data.json  — dữ liệu tổng hợp cho Remotion
    # ------------------------------------------------------------------
    # Copy background video sang vị trí Remotion expect
    bg_src  = os.path.join(ASSETS_DIR, "background_loop.mp4")
    bg_dest = os.path.join("topik-video", "public", "assets", "background.mp4")
    if os.path.exists(bg_src):
        os.makedirs(os.path.dirname(bg_dest), exist_ok=True)
        shutil.copy(bg_src, bg_dest)

    # Xây dựng payload JSON cho Remotion
    # Audio paths chuyển về relative (Remotion serve từ /public)
    relative_audio_paths = {
        key: f"/assets/{os.path.basename(path)}"
        for key, path in audio_paths.items()
    }

    # Merge audio_data timing into tiktok_script for Remotion
    tiktok_script = data_p3.get("tiktok_script", {})
    
    # Enrich video_1_news with segment timing
    if "video_1_news" in audio_data and audio_data["video_1_news"].get("segments"):
        tiktok_script["video_1_news"]["segments"] = audio_data["video_1_news"]["segments"]
        tiktok_script["video_1_news"]["total_duration"] = audio_data["video_1_news"]["total_duration"]
    
    # Enrich video_2_outline with part timing
    if "video_2_outline" in audio_data and audio_data["video_2_outline"].get("parts"):
        tiktok_script["video_2_outline"]["parts"] = audio_data["video_2_outline"]["parts"]
        tiktok_script["video_2_outline"]["total_duration"] = audio_data["video_2_outline"]["total_duration"]
    
    # Enrich quiz videos with split audio timing
    for quiz_key in ["video_3_vocab_quiz", "video_4_grammar_quiz"]:
        if quiz_key in audio_data:
            tiktok_script[quiz_key]["audio_timing"] = {
                "question": audio_data[quiz_key].get("question_audio"),
                "answer": audio_data[quiz_key].get("answer_audio"),
                "silence_duration": audio_data[quiz_key].get("silence_duration", 4.0),
                "total_duration": audio_data[quiz_key].get("total_duration", 0)
            }

    # Build final_data JSON
    final_data = {
        "meta":           data_p3.get("meta", {}),
        "phase1":         data_p1,
        "phase2":         data_p2,
        "phase4":         data_p4 if include_deep_dive else {},   # Deep Dive data
        "tiktok_script":  tiktok_script,              # Enriched with timing
        "word_doc_data":  data_p3.get("word_doc_data", {}),
        "audio_paths":    relative_audio_paths,       # Relative paths cho Remotion
        "audio_data":     audio_data,                 # Full timing data
        "video_bg":       "/assets/background.mp4",
        "video_bg_duration": video_bg_duration,       # NEW: Actual video duration in seconds
    }
    
    # Log video background info
    if video_bg_duration > 0:
        logging.info(f"📹 Video background duration: {video_bg_duration:.2f}s")
    else:
        logging.warning("⚠️ Video background duration unknown, Remotion will use default fallback")
    
    # Add video_5_deep_dive to tiktok_script if available
    if include_deep_dive and data_p4.get("video_5_deep_dive"):
        final_data["tiktok_script"]["video_5_deep_dive"] = data_p4["video_5_deep_dive"]
        # Enrich with audio timing data
        if "video_5_deep_dive" in audio_data:
            final_data["tiktok_script"]["video_5_deep_dive"]["segments"] = audio_data["video_5_deep_dive"].get("segments", [])
            final_data["tiktok_script"]["video_5_deep_dive"]["total_duration"] = audio_data["video_5_deep_dive"].get("total_duration", 0)
            final_data["tiktok_script"]["video_5_deep_dive"]["timestamps"] = audio_data["video_5_deep_dive"].get("timestamps", [])

    json_path = os.path.join(OUTPUT_DIR, "final_data.json")
    with open(json_path, "w", encoding="utf-8") as f:
        json.dump(final_data, f, ensure_ascii=False, indent=2)
    logging.info(f"💾 Đã lưu: {json_path}")

    # ------------------------------------------------------------------
    # GENERATE YOUTUBE METADATA (Timestamps, Title, Hashtags)
    # ------------------------------------------------------------------
    if include_deep_dive:
        youtube_info_path = os.path.join(OUTPUT_DIR, "youtube_info.txt")
        generate_youtube_description(final_data, youtube_info_path)
    
    # ------------------------------------------------------------------
    # RENDER LOOP — 5 video (4 TikTok + 1 YouTube Deep Dive)
    # ------------------------------------------------------------------
    rendered_videos = render_all_videos(json_path, include_deep_dive=include_deep_dive)

    # ------------------------------------------------------------------
    # UPLOAD LÊN GOOGLE DRIVE
    # ------------------------------------------------------------------
    DRIVE_FOLDER_ID = os.getenv("DRIVE_FOLDER_ID")
    ENABLE_YOUTUBE_UPLOAD = os.getenv("ENABLE_YOUTUBE_UPLOAD", "false").lower() == "true"
    YOUTUBE_PRIVACY = os.getenv("YOUTUBE_PRIVACY", "unlisted")  # public, unlisted, private
    YOUTUBE_PLAYLIST_ID = os.getenv("YOUTUBE_PLAYLIST_ID", "")

    # --- Google Drive Upload ---
    if DRIVE_FOLDER_ID:
        logging.info("-" * 60)
        logging.info("☁️  Bắt đầu Upload lên Google Drive...")

        # --- Upload Word ---
        if docx_path and os.path.exists(docx_path):
            upload_to_drive(docx_path, DRIVE_FOLDER_ID)
        else:
            logging.warning("⚠️  Không tìm thấy file Word để upload.")

        # --- Upload YouTube metadata (if Deep Dive was generated) ---
        if youtube_info_path and os.path.exists(youtube_info_path):
            logging.info("📝 Upload YouTube metadata...")
            upload_to_drive(youtube_info_path, DRIVE_FOLDER_ID)
        
        # --- Upload các video đã render thành công ---
        if rendered_videos:
            for vid_path in rendered_videos:
                if os.path.exists(vid_path) and os.path.getsize(vid_path) > 1024 * 1024:
                    logging.info(f"🎬 Upload Drive: {os.path.basename(vid_path)}")
                    file_id = upload_to_drive(vid_path, DRIVE_FOLDER_ID)
                    if file_id:
                        logging.info(f"   ✅ Drive Upload OK — ID: {file_id}")
                    else:
                        logging.error(f"   ❌ Drive Upload thất bại: {vid_path}")
                else:
                    logging.warning(f"⚠️  Bỏ qua file nhỏ hoặc không tồn tại: {vid_path}")
        else:
            # Fallback: quét toàn thư mục như logic gốc (an toàn)
            logging.warning("⚠️  Render loop không tạo video — thử quét thư mục...")
            for root, _dirs, files in os.walk("."):
                for fname in files:
                    if fname.endswith(".mp4") and "background" not in fname:
                        full = os.path.join(root, fname)
                        if os.path.getsize(full) > 1024 * 1024:
                            logging.info(f"🎬 Tìm thấy video rogue: {full}")
                            upload_to_drive(full, DRIVE_FOLDER_ID)
    else:
        logging.warning("⚠️ Thiếu DRIVE_FOLDER_ID — bỏ qua Drive upload.")

    # ------------------------------------------------------------------
    # UPLOAD LÊN YOUTUBE (Tùy chọn)
    # ------------------------------------------------------------------
    youtube_results = []
    
    if ENABLE_YOUTUBE_UPLOAD and YOUTUBE_UPLOAD_AVAILABLE and rendered_videos:
        logging.info("-" * 60)
        logging.info("📺 Bắt đầu Upload lên YouTube...")
        
        try:
            youtube_uploader = YouTubeUploader()
            if youtube_uploader.authenticate():
                # Get channel info
                channel_info = youtube_uploader.get_channel_info()
                if channel_info:
                    logging.info(f"   📺 Channel: {channel_info['title']}")
                
                # Phân loại video
                tiktok_videos = [v for v in rendered_videos if "V5_DeepDive" not in v]
                deep_dive_videos = [v for v in rendered_videos if "V5_DeepDive" in v]
                
                # Upload TikTok videos as Shorts
                if tiktok_videos:
                    logging.info(f"   🎬 Uploading {len(tiktok_videos)} TikTok Shorts...")
                    shorts_results = upload_tiktok_to_youtube(
                        video_paths=tiktok_videos,
                        video_data=final_data,
                        uploader=youtube_uploader,
                        playlist_id=YOUTUBE_PLAYLIST_ID if YOUTUBE_PLAYLIST_ID else None,
                        privacy=YOUTUBE_PRIVACY
                    )
                    youtube_results.extend(shorts_results)
                    
                    successful = [r for r in shorts_results if r.get("success")]
                    logging.info(f"   ✅ Shorts: {len(successful)}/{len(tiktok_videos)} uploaded")
                
                # Upload Deep Dive video (long-form)
                if deep_dive_videos:
                    logging.info("   🎥 Uploading Deep Dive video...")
                    for deep_video in deep_dive_videos:
                        dd_result = upload_deep_dive_to_youtube(
                            video_path=deep_video,
                            video_data=final_data,
                            youtube_info_path=youtube_info_path,
                            uploader=youtube_uploader,
                            privacy=YOUTUBE_PRIVACY
                        )
                        youtube_results.append(dd_result)
                        
                        if dd_result.get("success"):
                            logging.info(f"   ✅ Deep Dive uploaded: {dd_result.get('url')}")
                        else:
                            logging.error(f"   ❌ Deep Dive upload failed: {dd_result.get('error')}")
            else:
                logging.error("❌ YouTube authentication failed!")
                
        except Exception as e:
            logging.error(f"❌ YouTube upload error: {e}")
            traceback.print_exc()
    
    elif ENABLE_YOUTUBE_UPLOAD and not YOUTUBE_UPLOAD_AVAILABLE:
        logging.warning("⚠️ YouTube upload enabled but youtube_uploader module not available.")
    
    elif not ENABLE_YOUTUBE_UPLOAD:
        logging.info("ℹ️  YouTube upload disabled (set ENABLE_YOUTUBE_UPLOAD=true to enable)")

    # ------------------------------------------------------------------
    # GENERATE BLOG (Tùy chọn)
    # ------------------------------------------------------------------
    blog_result = None
    ENABLE_BLOG = os.getenv("ENABLE_BLOG", "true").lower() == "true"
    
    if ENABLE_BLOG and BLOG_GENERATOR_AVAILABLE:
        logging.info("-" * 60)
        logging.info("📝 Generating Blog Post...")
        
        try:
            blog_result = generate_blog_from_data(json_path, "blog_output")
            if blog_result:
                logging.info(f"   ✅ Blog generated: {blog_result.get('slug')}")
        except Exception as e:
            logging.error(f"❌ Blog generation error: {e}")
            traceback.print_exc()
    
    # ------------------------------------------------------------------
    # GENERATE PODCAST (Tùy chọn)
    # ------------------------------------------------------------------
    podcast_result = None
    ENABLE_PODCAST = os.getenv("ENABLE_PODCAST", "true").lower() == "true"
    
    if ENABLE_PODCAST and PODCAST_GENERATOR_AVAILABLE:
        logging.info("-" * 60)
        logging.info("🎙️ Generating Podcast Episode...")
        
        try:
            assets_dir = os.path.join(os.path.dirname(json_path), "assets")
            # Calculate episode number from date
            episode_num = int(datetime.now().strftime("%j"))  # Day of year
            
            podcast_result = generate_podcast_from_data(
                json_path, 
                assets_dir, 
                "podcast_output",
                episode_num
            )
            if podcast_result:
                logging.info(f"   ✅ Podcast generated: {podcast_result.get('filename')} ({podcast_result.get('duration_str')})")
        except Exception as e:
            logging.error(f"❌ Podcast generation error: {e}")
            traceback.print_exc()
    
    # ------------------------------------------------------------------
    # DEPLOY BLOG TO GITHUB PAGES (Tùy chọn)
    # ------------------------------------------------------------------
    ENABLE_GITHUB_DEPLOY = os.getenv("ENABLE_GITHUB_DEPLOY", "false").lower() == "true"
    
    if ENABLE_GITHUB_DEPLOY and GITHUB_DEPLOYER_AVAILABLE and blog_result:
        logging.info("-" * 60)
        logging.info("🚀 Deploying Blog to GitHub Pages...")
        
        try:
            deploy_success = deploy_blog_to_github("blog_output")
            if deploy_success:
                logging.info("   ✅ Blog deployed to GitHub Pages!")
            else:
                logging.error("   ❌ GitHub deployment failed")
        except Exception as e:
            logging.error(f"❌ GitHub deploy error: {e}")
            traceback.print_exc()
    
    # ------------------------------------------------------------------
    # PUBLISH TO SOCIAL MEDIA (Tùy chọn)
    # ------------------------------------------------------------------
    social_results = {}
    ENABLE_SOCIAL_MEDIA = os.getenv("ENABLE_SOCIAL_MEDIA", "false").lower() == "true"
    
    if ENABLE_SOCIAL_MEDIA and SOCIAL_PUBLISHER_AVAILABLE:
        logging.info("-" * 60)
        logging.info("📱 Publishing to Social Media...")
        
        try:
            social_results = publish_to_social_media(json_path)
            logging.info(f"   📱 Twitter: {'✅' if social_results.get('twitter') else '❌'}")
            logging.info(f"   📱 Telegram: {'✅' if social_results.get('telegram') else '❌'}")
            logging.info(f"   📱 Discord: {'✅' if social_results.get('discord') else '❌'}")
            logging.info(f"   📧 Email: {social_results.get('email', 0)} sent")
        except Exception as e:
            logging.error(f"❌ Social media error: {e}")
            traceback.print_exc()

    # ------------------------------------------------------------------
    # MONETIZATION: Generate Digital Products (Tùy chọn)
    # ------------------------------------------------------------------
    monetization_results = {}
    ENABLE_MONETIZATION = os.getenv("ENABLE_MONETIZATION", "true").lower() == "true"
    
    if ENABLE_MONETIZATION and MONETIZATION_AVAILABLE:
        logging.info("-" * 60)
        logging.info("💰 Generating Monetization Assets...")
        
        try:
            monetization_manager = MonetizationManager()
            monetization_results = monetization_manager.process_daily(final_data)
            
            if monetization_results.get("anki_deck"):
                logging.info(f"   📚 Anki Deck: {monetization_results['anki_deck']}")
            if monetization_results.get("lead_magnet"):
                logging.info(f"   📄 Lead Magnet PDF: {monetization_results['lead_magnet']}")
            if monetization_results.get("premium_content"):
                logging.info(f"   ⭐ Premium Content: {monetization_results['premium_content']}")
            
            # Upload products to Drive for distribution
            if DRIVE_FOLDER_ID:
                for key in ["anki_deck", "lead_magnet", "premium_content"]:
                    file_path = monetization_results.get(key)
                    if file_path and os.path.exists(file_path):
                        upload_to_drive(file_path, DRIVE_FOLDER_ID)
                        
        except Exception as e:
            logging.error(f"❌ Monetization error: {e}")
            traceback.print_exc()

    # ------------------------------------------------------------------
    # TELEGRAM BOT: Send Daily Push (Tùy chọn)
    # ------------------------------------------------------------------
    ENABLE_TELEGRAM_PUSH = os.getenv("ENABLE_TELEGRAM_PUSH", "false").lower() == "true"
    TELEGRAM_CHANNEL_ID = os.getenv("TELEGRAM_CHANNEL_ID", "")
    TELEGRAM_BOT_TOKEN = os.getenv("TELEGRAM_BOT_TOKEN", "")
    
    if ENABLE_TELEGRAM_PUSH and TELEGRAM_BOT_AVAILABLE and TELEGRAM_CHANNEL_ID:
        logging.info("-" * 60)
        logging.info("🤖 Sending Telegram Daily Push...")
        
        try:
            asyncio.run(send_daily_push(TELEGRAM_BOT_TOKEN, TELEGRAM_CHANNEL_ID, json_path))
            logging.info("   ✅ Telegram push sent!")
        except Exception as e:
            logging.error(f"❌ Telegram push error: {e}")
            traceback.print_exc()

    # --- Summary ---
    logging.info("=" * 60)
    logging.info("🏁 HOÀN THÀNH — Toàn bộ pipeline đã chạy xong.")
    logging.info(f"   📄 Word: {docx_path or 'N/A'}")
    logging.info(f"   🎬 Videos rendered: {len(rendered_videos)} / 5")
    if DRIVE_FOLDER_ID:
        logging.info(f"   ☁️  Drive: Uploaded to folder {DRIVE_FOLDER_ID}")
    if youtube_results:
        successful_yt = [r for r in youtube_results if r.get("success")]
        logging.info(f"   📺 YouTube: {len(successful_yt)}/{len(youtube_results)} uploaded")
        for r in successful_yt:
            logging.info(f"      → {r.get('url')} ({r.get('privacy')})")
    if blog_result:
        logging.info(f"   📝 Blog: {blog_result.get('slug')}")
    if podcast_result:
        logging.info(f"   🎙️ Podcast: {podcast_result.get('filename')}")
    if social_results:
        logging.info(f"   📱 Social: Twitter={social_results.get('twitter')}, Telegram={social_results.get('telegram')}")
    if monetization_results:
        logging.info(f"   💰 Products: Anki={bool(monetization_results.get('anki_deck'))}, PDF={bool(monetization_results.get('lead_magnet'))}")
    if include_deep_dive:
        logging.info(f"   📝 YouTube metadata: {youtube_info_path or 'N/A'}")
    logging.info("=" * 60)


# ==============================================================================
if __name__ == "__main__":
    main()
